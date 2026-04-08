"""
上传与解析接口 - 负责人: 成员1（路由层）
解析逻辑由成员2实现: modules/parser/document_parser.py
"""
import uuid
import os
import re
import hashlib
import asyncio
from datetime import datetime, timezone
from typing import Literal, Optional

from fastapi import APIRouter, UploadFile, File, Depends
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session
 
from config import settings
from db.database import get_db
from db.models import FileRecord, TaskRecord
from errors import AppError
 
router = APIRouter(tags=["文档上传与解析"])
 
ALLOWED_MIME = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}
 
 
def _safe_filename(filename: str) -> str:
    """去除路径符号，只保留安全文件名"""
    return os.path.basename(filename).replace("..", "").strip()
 
 
def _file_md5(content: bytes) -> str:
    return hashlib.md5(content).hexdigest()
 
 
# ── POST /upload ──────────────────────────────────────────────
@router.post("/upload", summary="上传文件")
async def upload_file(
    file: UploadFile = File(..., description="支持 PDF / DOCX / XLSX"),
    db: Session = Depends(get_db),
):
    """
    上传文件，返回 file_id。
    - 自动校验格式与大小
    - 相同内容的文件不重复存储（MD5去重）
    """
    # ── 文件名安全处理 ──
    safe_name = _safe_filename(file.filename or "unnamed")
    ext = safe_name.rsplit(".", 1)[-1].lower() if "." in safe_name else ""
 
    if ext not in settings.allowed_extensions:
        raise AppError(400, "INVALID_FILE_TYPE", f"不支持的文件格式: .{ext}，允许: {settings.allowed_extensions}")
 
    # ── 读取内容 & 大小校验 ──
    content = await file.read()
    size_mb = len(content) / (1024 * 1024)
    if size_mb > settings.max_file_size_mb:
        raise AppError(400, "FILE_TOO_LARGE", f"文件 {size_mb:.1f}MB 超过 {settings.max_file_size_mb}MB 限制")
 
    # ── MD5去重：相同文件直接返回已有 file_id ──
    md5 = _file_md5(content)
    existing = db.query(FileRecord).filter_by(md5=md5).first()
    if existing:
        return {
            "file_id": existing.file_id,
            "filename": existing.filename,
            "status": existing.status,
            "duplicate": True,
        }
 
    # ── 保存文件 ──
    file_id = str(uuid.uuid4())
    os.makedirs(settings.upload_dir, exist_ok=True)
    file_path = os.path.join(settings.upload_dir, f"{file_id}.{ext}")
    with open(file_path, "wb") as f:
        f.write(content)
 
    # ── 写入数据库 ──
    now = datetime.now(timezone.utc).isoformat()
    record = FileRecord(
        file_id=file_id,
        filename=safe_name,
        file_type=ext,
        file_path=file_path,
        file_size=len(content),
        md5=md5,
        status="uploaded",
        uploaded_at=now,
    )
    db.add(record)
    db.commit()
 
    return {
        "file_id": file_id,
        "filename": safe_name,
        "size_mb": round(size_mb, 2),
        "status": "uploaded",
        "duplicate": False,
    }
 
 
# ── POST /parse ───────────────────────────────────────────────
class ParseRequest(BaseModel):
    file_id: str
    priority: int = 1  # 预留优先级字段，默认普通优先级
 
 
@router.post("/parse", summary="提交解析任务")
async def parse_file(body: ParseRequest, db: Session = Depends(get_db)):
    """
    提交解析任务（异步），立即返回 task_id。
    前端每2秒轮询 GET /parse/status/{task_id}。
    """
    record = db.query(FileRecord).filter_by(file_id=body.file_id).first()
    if not record:
        raise AppError(404, "FILE_NOT_FOUND", f"文件 {body.file_id} 不存在")
 
    # 防止重复提交：同一文件已有 pending/processing 任务时直接返回
    existing_task = (
        db.query(TaskRecord)
        .filter_by(file_id=body.file_id, task_type="parse")
        .filter(TaskRecord.status.in_(["pending", "processing"]))
        .first()
    )
    if existing_task:
        return {
            "task_id": existing_task.task_id,
            "file_id": body.file_id,
            "status": existing_task.status,
            "message": "该文件已有进行中的解析任务",
        }
 
    # ── 创建任务记录 ──
    task_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    task = TaskRecord(
        task_id=task_id,
        file_id=body.file_id,
        task_type="parse",
        status="pending",
        progress=0,
        created_at=now,
        updated_at=now,
    )
    db.add(task)
    db.commit()
 
    # ── 后台异步执行 ──
    asyncio.create_task(_run_parse(task_id, body.file_id, record.file_path))
 
    return {"task_id": task_id, "file_id": body.file_id, "status": "pending"}
 
 
# ── GET /parse/status/{task_id} ───────────────────────────────
class ParseStatusResponse(BaseModel):
    """解析任务状态（轮询接口返回体）"""

    task_id: str = Field(..., description="任务 ID")
    file_id: str = Field(..., description="关联文件 ID")
    status: Literal["pending", "processing", "done", "failed"] = Field(
        ..., description="任务状态"
    )
    progress: int = Field(..., ge=0, le=100, description="进度 0~100")
    updated_at: str = Field(..., description="最后更新时间 ISO8601")
    chunk_count: Optional[int] = Field(
        None, description="解析完成后的分块数量（仅 status=done 时有）"
    )
    error: Optional[str] = Field(
        None, description="失败原因（仅 status=failed 时有）"
    )

    model_config = {"json_schema_extra": {"example": {
        "task_id": "7cafc2c9-36be-4481-8064-2295fb450e96",
        "file_id": "1a635ee3-f56a-46d4-905a-bf742287c5b1",
        "status": "done",
        "progress": 100,
        "updated_at": "2026-03-24T16:22:56.029250+00:00",
        "chunk_count": 3,
    }}}


@router.get(
    "/parse/status/{task_id}",
    summary="查询解析任务状态",
    response_model=ParseStatusResponse,
    response_model_exclude_none=True,
)
async def parse_status(task_id: str, db: Session = Depends(get_db)) -> ParseStatusResponse:
    """
    轮询解析任务状态，建议每2秒一次。
    返回字段:
    - status: pending | processing | done | failed
    - progress: 0~100
    - chunk_count: 解析完成后的分块数量
    - error: 失败原因（仅 failed 时有）
    """
    task = db.query(TaskRecord).filter_by(task_id=task_id).first()
    if not task:
        raise AppError(404, "TASK_NOT_FOUND", f"任务 {task_id} 不存在")
 
    base = {
        "task_id": task_id,
        "file_id": task.file_id,
        "status": task.status,
        "progress": task.progress,
        "updated_at": task.updated_at,
    }
 
    if task.status == "done":
        file_record = db.query(FileRecord).filter_by(file_id=task.file_id).first()
        base["chunk_count"] = file_record.chunk_count if file_record else 0
    elif task.status == "failed":
        base["error"] = task.error_msg
 
    return ParseStatusResponse(**base)
 
 
# ── 后台解析任务 ──────────────────────────────────────────────
async def _run_parse(task_id: str, file_id: str, file_path: str):
    """
    后台任务调度：
      10% → 开始解析（成员2）
      70% → 解析完成，开始建索引（成员3）
      90% → 索引完成
     100% → done
    """
    from db.database import SessionLocal
    db = SessionLocal()
    try:
        _update_task(db, task_id, status="processing", progress=5)

        # ── 进度回调：parse_document 内部会调用此回调回传细粒度进度 ──
        def _on_parse_progress(pct: int, msg: str):
            # pct 0-70 → 映射到数据库 10-49
            mapped = min(49, max(10, pct))
            _update_task(db, task_id, status="processing", progress=mapped)

        # ── 调用成员2：文档解析 ──────────────────────────────
        from modules.parser.document_parser import parse_document
        result = await asyncio.get_event_loop().run_in_executor(
            None, parse_document, file_path, file_id, _on_parse_progress
        )
        # ────────────────────────────────────────────────────
 
        if not result or not result.get("chunks"):
            raise ValueError("解析结果为空，chunk 列表不能为空")

        # 解析完成立即更新 50%，避免前端进度长时间停在 10% 不动
        _update_task(db, task_id, status="processing", progress=50)

        # ── 调用成员3：建立检索索引 ──────────────────────────
        from modules.retriever.indexer import build_index

        def _on_index_progress(pct: int, msg: str):
            mapped = min(89, max(70, pct))
            _update_task(db, task_id, status="processing", progress=mapped)

        # 必须为 True：同一 file_id 重新解析时若跳过建索引，Chroma/BM25 仍保留旧正文，
        # 会出现「界面是排名 PDF、检索却命中量子计算测试文」等严重串档。
        ok = await asyncio.get_event_loop().run_in_executor(
            None, build_index, result, True, _on_index_progress
        )
        # ────────────────────────────────────────────────────
        if not ok:
            raise RuntimeError("向量索引构建失败（常见原因：Ollama 未运行、未拉取 embedding 模型、或 /v1/embeddings 请求格式错误）")

        _update_task(db, task_id, status="processing", progress=90)
 
        # ── 更新文件状态 ──
        record = db.query(FileRecord).filter_by(file_id=file_id).first()
        if record:
            now = datetime.now(timezone.utc).isoformat()
            record.status = "indexed"
            record.chunk_count = len(result.get("chunks", []))
            record.parsed_at = now
            record.indexed_at = now
            db.commit()

        try:
            from modules.cache.redis_client import invalidate_cache
            invalidate_cache("*")
        except Exception:
            pass

        _update_task(db, task_id, status="done", progress=100)
 
    except Exception as e:
        _update_task(db, task_id, status="failed", error_msg=str(e))
    finally:
        db.close()
 
 
def _update_task(db, task_id: str, status: str, progress: int = None, error_msg: str = None):
    """原子更新任务状态"""
    task = db.query(TaskRecord).filter_by(task_id=task_id).first()
    if task:
        task.status = status
        task.updated_at = datetime.now(timezone.utc).isoformat()
        if progress is not None:
            task.progress = progress
        if error_msg is not None:
            task.error_msg = error_msg
        db.commit()


# ── POST /parse/batch ─────────────────────────────────────────
@router.post("/parse/batch", summary="批量提交解析任务")
async def parse_batch(db: Session = Depends(get_db)):
    """
    一键解析所有 status=uploaded 的文件。
    返回每个文件的 task_id，前端可逐一轮询 /parse/status/{task_id}。
    """
    pending_files = db.query(FileRecord).filter_by(status="uploaded").all()
    if not pending_files:
        return {"message": "没有待解析的文件", "tasks": []}

    tasks = []
    for record in pending_files:
        # 跳过已有进行中任务的文件
        existing = (
            db.query(TaskRecord)
            .filter_by(file_id=record.file_id, task_type="parse")
            .filter(TaskRecord.status.in_(["pending", "processing"]))
            .first()
        )
        if existing:
            tasks.append({
                "file_id": record.file_id,
                "filename": record.filename,
                "task_id": existing.task_id,
                "status": "already_running",
            })
            continue

        task_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()
        task = TaskRecord(
            task_id=task_id,
            file_id=record.file_id,
            task_type="parse",
            status="pending",
            progress=0,
            created_at=now,
            updated_at=now,
        )
        db.add(task)
        db.commit()
        asyncio.create_task(_run_parse(task_id, record.file_id, record.file_path))
        tasks.append({
            "file_id": record.file_id,
            "filename": record.filename,
            "task_id": task_id,
            "status": "submitted",
        })

    return {
        "message": f"已提交 {len([t for t in tasks if t['status']=='submitted'])} 个解析任务",
        "tasks": tasks,
    }



@router.post("/template/placeholders", summary="解析模板占位符")
async def extract_placeholders(body: dict, db: Session = Depends(get_db)):
    """
    智能解析模板内容，自动识别要填写的字段。
    - 支持 {{占位符}} 格式（有明确标记）
    - 对于普通表格/文档，使用 LLM 智能推断需要填写的字段
    请求体: { "template_file_id": "xxx" }
    返回:  { "fields": ["合同金额", "甲方名称", ...], "method": "placeholder|llm" }
    """
    file_id = body.get("template_file_id")
    if not file_id:
        raise AppError(400, "INVALID_REQUEST", "缺少 template_file_id")

    record = db.query(FileRecord).filter_by(file_id=file_id).first()
    if not record:
        raise AppError(404, "FILE_NOT_FOUND", f"模板文件 {file_id} 不存在")

    if record.file_type not in ("docx", "xlsx"):
        raise AppError(400, "INVALID_FILE_TYPE", "仅支持 .docx 和 .xlsx 模板")

    try:
        # 先找 {{占位符}}
        fields = _extract_placeholders(record.file_path, record.file_type)

        if fields:
            return {"fields": fields, "count": len(fields), "method": "placeholder"}

        # 没有占位符 → 调用 LLM 智能分析模板内容，推断要填的字段
        text_content = _extract_template_text(record.file_path, record.file_type)
        if not text_content:
            return {"fields": [], "count": 0, "method": "none", "message": "模板内容为空"}

        llm_fields = await _llm_extract_fields(text_content, record.file_type)
        return {"fields": llm_fields, "count": len(llm_fields), "method": "llm"}

    except Exception as e:
        raise AppError(500, "PARSE_FAILED", f"解析模板失败：{str(e)}")


def _extract_template_text(file_path: str, file_type: str) -> str:
    """提取模板文本内容（不含格式），用于 LLM 分析。严格限制总字符数。"""
    MAX_CHARS = 1500  # 约 1000 tokens，留足够空间给 prompt

    if file_type == "docx":
        from docx import Document
        doc = Document(file_path)
        parts = []
        total = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if total + len(text) + 1 > MAX_CHARS:
                    break
                parts.append(text)
                total += len(text) + 1
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    if total + len(row_text) + 1 > MAX_CHARS:
                        break
                    parts.append(row_text)
                    total += len(row_text) + 1
            if total >= MAX_CHARS:
                break
        return "\n".join(parts)

    elif file_type == "xlsx":
        # Excel 智能提取：只取前3行（通常含表头），字符数极少
        from openpyxl import load_workbook
        wb = load_workbook(file_path, data_only=True)
        parts = []
        total = 0
        for sheet in wb.worksheets:
            parts.append(f"[Sheet: {sheet.title}]")
            total += len(sheet.title) + 12
            rows_taken = 0
            for row in sheet.iter_rows(values_only=True):
                if rows_taken >= 3 or total >= MAX_CHARS:
                    break
                row_vals = [str(c) if c is not None else "" for c in row]
                row_text = " | ".join(v for v in row_vals if v)
                if row_text:
                    parts.append(row_text)
                    total += len(row_text) + 1
                rows_taken += 1
        wb.close()
        return "\n".join(parts)

    return ""


async def _llm_extract_fields(text_content: str, file_type: str) -> list[str]:
    """
    调用 LLM 智能分析模板内容，推断需要填写的字段。
    例如：看到"合同金额"、"甲方"、"签署日期"等表头，推断出字段名。
    """
    from openai import OpenAI
    from config import settings

    prompt = f"""你是一个文档分析专家。请分析以下{file_type}模板的内容，识别出需要填写/补充的字段名称。

要求：
1. 只输出字段名，每行一个，不要任何解释
2. 优先输出表格表头、空格/横线处的标题
3. 不要输出已填写了具体数值的行（只输出空缺的、需要填写的）
4. 字段名要简洁，如"合同金额"、"甲方名称"，不要超过15个字
5. 只输出需要填写的字段，不要输出文档标题、说明文字等

{file_type}内容：
{text_content[:3000]}
"""

    client = OpenAI(api_key=settings.llm_api_key, base_url=settings.llm_base_url)
    resp = client.chat.completions.create(
        model=settings.llm_model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=300,
    )
    raw = resp.choices[0].message.content or ""

    # 解析 LLM 输出，每行一个字段名
    fields = []
    for line in raw.split("\n"):
        line = line.strip().strip("•-*1234567890.、 ")
        if line and 1 < len(line) <= 20:
            fields.append(line)

    # 去重
    seen = set()
    unique = []
    for f in fields:
        if f not in seen:
            seen.add(f)
            unique.append(f)
    return unique


def _extract_placeholders(file_path: str, file_type: str) -> list[str]:
    """从 docx/xlsx 文件中提取所有 {{字段名}} 占位符"""
    seen = set()
    fields = []

    if file_type == "docx":
        from docx import Document
        doc = Document(file_path)
        # 遍历段落
        for para in doc.paragraphs:
            for m in re.finditer(r'\{\{(.+?)\}\}', para.text):
                key = m.group(1).strip()
                if key and key not in seen:
                    seen.add(key)
                    fields.append(key)
        # 遍历表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for m in re.finditer(r'\{\{(.+?)\}\}', para.text):
                            key = m.group(1).strip()
                            if key and key not in seen:
                                seen.add(key)
                                fields.append(key)

    elif file_type == "xlsx":
        from openpyxl import load_workbook
        wb = load_workbook(file_path, data_only=False)
        for sheet in wb.worksheets:
            merged_cells = {
                str(cell)
                for rng in sheet.merged_cells.ranges
                for cell in rng.cells
            }
            for row in sheet.iter_rows():
                for cell in row:
                    # 跳过合并单元格的从属格
                    if cell.coordinate in merged_cells:
                        is_master = any(
                            cell.row == rng.min_row and cell.column == rng.min_col
                            for rng in sheet.merged_cells.ranges
                        )
                        if not is_master:
                            continue
                    val = cell.value
                    if val and isinstance(val, str):
                        for m in re.finditer(r'\{\{(.+?)\}\}', val):
                            key = m.group(1).strip()
                            if key and key not in seen:
                                seen.add(key)
                                fields.append(key)
        wb.close()

    return fields