"""
表格回填模块 - 负责人: 成员4
函数签名已锁定，不得更改
"""
import re
import os
import shutil as _shutil

from loguru import logger
from docx import Document
from openpyxl import load_workbook


def _replace_placeholder_in_paragraph(para, field_map: dict) -> None:
    """
    替换段落中的占位符，处理占位符被拆分到多个 run 的情况。
    思路参考 SpreadsheetLLM CoS：先重建完整文本定位占位符，再按原 run 边界写回。
    """
    full_text = "".join(run.text for run in para.runs)
    new_text = full_text
    for key, val in field_map.items():
        new_text = new_text.replace(f"{{{{{key}}}}}", str(val))

    if new_text == full_text:
        return  # 无占位符，跳过

    # 将新文本按原 run 长度分配写回，保留各 run 的格式
    idx = 0
    for run in para.runs:
        run_len = len(run.text)
        run.text = new_text[idx: idx + run_len]
        idx += run_len
    # 若新文本比原文本长/短，剩余内容追加到最后一个 run
    if idx < len(new_text) and para.runs:
        para.runs[-1].text += new_text[idx:]


def fill_docx(template_path: str, answers: list, output_path: str) -> bool:
    """
    内部子函数：填充DOCX模板占位符
    :param template_path: 模板文件路径
    :param answers: 回填数据列表（fill_request['answers']）
    :param output_path: 输出文件路径
    :return: 填充成功返回True，失败返回False
    """
    try:
        doc = Document(template_path)
        field_map = {a['field_name']: a['value'] for a in answers}

        # 遍历顶层段落
        for para in doc.paragraphs:
            _replace_placeholder_in_paragraph(para, field_map)

        # 遍历文档表格（含表格内各段落，支持合并单元格）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_placeholder_in_paragraph(para, field_map)

        doc.save(output_path)
        logger.info(f"DOCX模板填充完成，输出文件：{output_path}")
        return True
    except Exception as e:
        logger.error(f"DOCX模板填充失败：{str(e)}")
        return False


def fill_xlsx(template_path: str, answers: list, output_path: str) -> bool:
    """
    内部子函数：填充XLSX模板占位符
    :param template_path: 模板文件路径
    :param answers: 回填数据列表（fill_request['answers']）
    :param output_path: 输出文件路径
    :return: 填充成功返回True，失败返回False
    """
    try:
        wb = load_workbook(template_path)
        field_map = {a['field_name']: a['value'] for a in answers}

        for sheet in wb.worksheets:
            # iter_rows 不含合并单元格的从属格，需单独处理合并区域
            merged_cells = {str(cell) for rng in sheet.merged_cells.ranges for cell in rng.cells}

            for row in sheet.iter_rows():
                for cell in row:
                    # 跳过合并单元格的从属格（只写主格，避免 openpyxl 抛异常）
                    if cell.coordinate in merged_cells:
                        is_master = any(
                            cell.row == rng.min_row and cell.column == rng.min_col
                            for rng in sheet.merged_cells.ranges
                        )
                        if not is_master:
                            continue
                    if cell.value and isinstance(cell.value, str):
                        new_val = cell.value
                        for key, val in field_map.items():
                            new_val = new_val.replace(f"{{{{{key}}}}}", str(val))
                        if new_val != cell.value:
                            cell.value = new_val

        wb.save(output_path)
        logger.info(f"XLSX模板填充完成，输出文件：{output_path}")
        return True
    except Exception as e:
        logger.error(f"XLSX模板填充失败：{str(e)}")
        return False


def fill_table(template_path: str, fill_request: dict, output_path: str, shutil=None) -> bool:
    """
    根据 FillRequest 将字段填入模板文件

    Args:
        template_path: 模板文件本地路径
        fill_request:  FillRequest dict（规范文档 4.4）
        output_path:   输出文件路径

    Returns:
        True 成功 / False 失败

    支持格式:
        .docx → 占位符格式 {{字段名}}
        .xlsx → 单元格内容为 {{字段名}}
    """
    logger.info(f"开始填表: template={template_path}, output={output_path}")

    # 1. 基础参数校验
    if not os.path.exists(template_path):
        logger.error(f"模板文件不存在：{template_path}")
        return False
    if not fill_request or not isinstance(fill_request.get('answers'), list):
        logger.error("FillRequest格式错误，answers必须为非空列表")
        return False
    if not output_path:
        logger.error("输出文件路径不能为空")
        return False

    # 2. 提取回填数据
    answers = fill_request['answers']

    # 3. 获取文件后缀，判断模板类型
    file_suffix = os.path.splitext(template_path)[-1].lstrip('.').lower()

    # 4. 按文件类型执行填充逻辑
    if file_suffix == 'docx':
        fill_success = fill_docx(template_path, answers, output_path)
    elif file_suffix == 'xlsx':
        fill_success = fill_xlsx(template_path, answers, output_path)
    else:
        logger.warning(f"不支持的文件格式：{file_suffix}，复制原文件到输出路径")
        try:
            _shutil.copy(template_path, output_path)
            fill_success = True
        except Exception as e:
            logger.error(f"复制文件失败: {e}")
            fill_success = False

    # 5. 返回最终结果
    if fill_success:
        logger.info(f"填表完成: output={output_path}")
    else:
        logger.error(f"填表失败: template={template_path}")
    return fill_success
