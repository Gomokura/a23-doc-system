"""
智能填表模块 - 自动从数据源文档中提取信息并填写模板
支持：无占位符模板（通过 LLM 识别表头）和有占位符模板（{{字段名}}）
"""
import os
import json
import re

from loguru import logger
from openai import OpenAI
from config import settings


# ────────────────────────────────────────────────────────────
# 工具函数：规则层行筛选（无需 LLM，直接 pandas）
# ────────────────────────────────────────────────────────────

def _rule_based_filter(df, instruction: str):
    """
    尝试用规则解析简单筛选条件，返回过滤后的 DataFrame。
    无法解析时返回 None，由调用方降级到 LLM。

    支持的模式：
      - "城市是青岛市" / "城市等于青岛市" / "城市为青岛市"
      - "城市包含青岛" / "城市含有青岛"
      - "金额大于1000" / "金额>1000" / "金额>=1000"
      - "金额小于500" / "金额<500" / "金额<=500"
      - "状态不是已取消" / "状态不等于已取消"
      - 多条件 AND：用"且"/"并且"/"，"连接
    """
    import pandas as pd

    instruction = instruction.strip()
    cols = list(df.columns)
    logger.debug(f"[规则筛选] 指令: {instruction!r}, 列名: {cols}")

    import pandas as pd

    def _find_col(keyword: str):
        """模糊匹配列名，去掉常见前缀词后再匹配"""
        keyword = re.sub(r'^(只要|仅|筛选|过滤|选择|保留)', '', keyword).strip()
        for c in cols:
            if keyword in str(c) or str(c) in keyword:
                return c
        return None

    # 含有排除/否定逻辑时，尝试规则层处理排除条件
    # 模式：先处理包含条件，再处理排除条件
    # 如果同时有包含+排除，分别处理后合并
    exclude_match = re.search(r'不要(.+?)(?:的|$)', instruction)
    include_match = re.search(r'(?:只要|只有|仅)(.+?)(?:的|不要|$)', instruction)

    if exclude_match or include_match:
        result = df.copy()

        # 先处理包含条件
        if include_match:
            include_str = include_match.group(1).strip()
            # 解析"X是Y"/"X为Y"/"X等于Y"格式
            inc_cond = re.search(r'(.{1,10}?)(?:是|为|等于)(.+)', include_str)
            if inc_cond:
                col_kw = inc_cond.group(1).strip()
                val = inc_cond.group(2).strip().rstrip('的')
                logger.debug(f"[规则筛选] include条件: col_kw={col_kw!r}, val={val!r}")
                # 精确找列名
                col = _find_col(col_kw)
                logger.debug(f"[规则筛选] _find_col({col_kw!r}) → {col!r}")
                if col:
                    matched = result[result[col].astype(str).str.strip() == val]
                    logger.debug(f"[规则筛选] 精确列匹配: {col}=={val!r} → {len(matched)} 行")
                    result = matched
                else:
                    # 找不到精确列，在所有列（含非object）里找值=='必修'的行
                    mask = pd.Series([False] * len(result), index=result.index)
                    for tc in result.columns:
                        try:
                            m2 = result[tc].astype(str).str.strip() == val
                            if m2.any():
                                logger.debug(f"[规则筛选] 在列 {tc!r} 找到 {val!r}，{m2.sum()} 行")
                            mask = mask | m2
                        except Exception:
                            pass
                    if mask.any():
                        result = result[mask]
                    else:
                        logger.debug(f"[规则筛选] 所有列均未找到值 {val!r}，include条件无效")
            else:
                # 没有"是/为/等于"，尝试直接匹配列值
                include_vals = re.split(r'[和或、,，]', include_str)
                include_vals = [v.strip().rstrip('的') for v in include_vals if v.strip()]
                for val in include_vals:
                    col = _find_col(val)
                    if col:
                        result = result[result[col].astype(str).str.strip() == val]

        # 再处理排除条件（在课程名/所有文本列中排除关键词）
        if exclude_match:
            exclude_str = exclude_match.group(1).strip()
            # 去掉末尾"的"
            exclude_str = exclude_str.rstrip('的').strip()
            exclude_vals = re.split(r'[和、,，]', exclude_str)
            exclude_vals = [v.strip().rstrip('类') if v.strip().endswith('类') and len(v.strip()) > 2 else v.strip()
                           for v in exclude_vals if v.strip()]
            logger.debug(f"[规则筛选] 排除关键词: {exclude_vals}")

            # 找文本列（优先课程名，其次所有object列）
            text_cols = [c for c in result.columns if result[c].dtype == object]
            for val in exclude_vals:
                # 先看能不能精确匹配某列名
                col = _find_col(val)
                if col:
                    result = result[result[col].astype(str).str.strip() != val]
                else:
                    # 在所有文本列中做完整关键词排除（要求val长度>=2，防止单字误杀）
                    if len(val) < 2:
                        continue
                    mask = pd.Series([False] * len(result), index=result.index)
                    for tc in text_cols:
                        mask = mask | result[tc].astype(str).str.contains(val, na=False, regex=False)
                    result = result[~mask]

        if len(result) < len(df):  # 确实有过滤效果才返回
            logger.debug(f"[规则筛选] 包含/排除规则命中，{len(df)} → {len(result)} 行")
            return result.reset_index(drop=True)

    # 含有复杂排除逻辑且规则层无法处理时，降级到 LLM
    if re.search(r'不要|排除|除了|除去|去掉|不包含|不含', instruction):
        logger.debug("[规则筛选] 检测到排除条件但规则层未命中，降级到 LLM")
        return None

    def _apply_single(df_, cond: str):
        """解析并应用单个条件，失败返回 None"""
        cond = cond.strip()

        # 日期范围：X在A到B之间 / X从A到B / A<=X<=B
        m = re.search(r'(.{1,10}?)(?:在|从)(.{4,20}?)(?:到|至|~)(.{4,20})(?:之间|期间)?', cond)
        if m:
            col = _find_col(m.group(1).strip())
            d1 = m.group(2).strip().strip('"\'""')
            d2 = m.group(3).strip().strip('"\'""')
            if col is not None:
                try:
                    s = pd.to_datetime(d1, errors='coerce')
                    e = pd.to_datetime(d2, errors='coerce')
                    col_dt = pd.to_datetime(df_[col], errors='coerce')
                    if pd.notna(s) and pd.notna(e):
                        return df_[(col_dt >= s) & (col_dt <= e)]
                except Exception:
                    pass

        # 等于：X是Y / X等于Y / X为Y / X=Y
        m = re.search(r'(.{1,10}?)(?:是|等于|为|=)(.+)', cond)
        if m:
            col = _find_col(m.group(1).strip())
            val = m.group(2).strip().strip('"\'""').rstrip('的')
            if col is not None:
                return df_[df_[col].astype(str).str.strip() == val]

        # 包含：X包含Y / X含有Y / X有Y
        m = re.search(r'(.{1,10}?)(?:包含|含有|含)(.+)', cond)
        if m:
            col = _find_col(m.group(1).strip())
            val = m.group(2).strip().strip('"\'""').rstrip('的')
            if col is not None:
                return df_[df_[col].astype(str).str.contains(val, na=False)]

        # 不等于：X不是Y / X不等于Y / X!=Y
        m = re.search(r'(.{1,10}?)(?:不是|不等于|!=)(.+)', cond)
        if m:
            col = _find_col(m.group(1).strip())
            val = m.group(2).strip().strip('"\'""').rstrip('的')
            if col is not None:
                return df_[df_[col].astype(str).str.strip() != val]

        # 数值比较：X大于/>=/>Y  X小于/<=/< Y
        m = re.search(r'(.{1,10}?)(?:大于等于|>=|≥)(\d+\.?\d*)', cond)
        if m:
            col = _find_col(m.group(1).strip())
            val = float(m.group(2))
            if col is not None:
                return df_[pd.to_numeric(df_[col], errors='coerce') >= val]

        m = re.search(r'(.{1,10}?)(?:大于|>)(\d+\.?\d*)', cond)
        if m:
            col = _find_col(m.group(1).strip())
            val = float(m.group(2))
            if col is not None:
                return df_[pd.to_numeric(df_[col], errors='coerce') > val]

        m = re.search(r'(.{1,10}?)(?:小于等于|<=|≤)(\d+\.?\d*)', cond)
        if m:
            col = _find_col(m.group(1).strip())
            val = float(m.group(2))
            if col is not None:
                return df_[pd.to_numeric(df_[col], errors='coerce') <= val]

        m = re.search(r'(.{1,10}?)(?:小于|<)(\d+\.?\d*)', cond)
        if m:
            col = _find_col(m.group(1).strip())
            val = float(m.group(2))
            if col is not None:
                return df_[pd.to_numeric(df_[col], errors='coerce') < val]

        return None

    # 拆分多条件（且/并且/，）
    parts = re.split(r'[，,]|且|并且', instruction)
    parts = [p.strip() for p in parts if p.strip()]

    result = df.copy()
    matched_any = False
    for part in parts:
        filtered = _apply_single(result, part)
        if filtered is not None:
            result = filtered
            matched_any = True
        else:
            # 有一个条件解析失败，整体降级到 LLM
            return None

    return result.reset_index(drop=True) if matched_any else None


# ────────────────────────────────────────────────────────────
# 工具函数：LLM 通用行筛选（任意列、任意条件）
# ────────────────────────────────────────────────────────────

def _parse_filter_intent(user_instruction: str, columns: list) -> list | None:
    """
    用 LLM 将自然语言筛选指令解析为结构化过滤条件（一次调用，不处理数据行）。

    返回格式（OR 分组，每组内部为 AND）：
    {
      "logic": "or",
      "groups": [
        [{"col": "监测时间", "op": "eq", "v1": "2025-11-25 09:00:00.0"}, {"col": "城市", "op": "eq", "v1": "德州市"}],
        [{"col": "监测时间", "op": "eq", "v1": "2025-11-25 09:00:00.0"}, {"col": "城市", "op": "eq", "v1": "潍坊市"}],
      ]
    }
    也兼容旧格式（纯列表，视为单组 AND）。
    如果指令不是筛选条件，返回 None。
    """
    client = OpenAI(
        api_key=settings.llm_api_key,
        base_url=settings.llm_base_url,
        default_headers=settings.openai_default_headers,
        timeout=60,
    )

    cols_str = "、".join(columns)
    prompt = f"""你是数据筛选条件解析助手。请将用户的自然语言筛选指令解析为结构化 JSON。

可用列名：{cols_str}

支持的操作符（op）：
- eq：等于  - neq：不等于  - contains：包含  - not_contains：不包含
- gt：大于  - lt：小于  - gte：大于等于  - lte：小于等于
- between：区间（需要 v1 开始值、v2 结束值）
- empty：为空  - not_empty：不为空

【重要】如果指令中包含多个并列的条件组（如"表一城市A，表二城市B，表三城市C"），
每个组之间是 OR 关系（取并集），组内条件是 AND 关系（取交集）。

输出格式：
{{
  "logic": "or",
  "groups": [
    [{{"col": "列名", "op": "操作符", "v1": "值"}}],
    [{{"col": "列名", "op": "操作符", "v1": "值"}}]
  ]
}}

如果只有一组条件（AND关系），groups 只有一个元素。
如果指令不是筛选条件（如纯填表说明、模板格式描述等），输出 null。

用户指令：{user_instruction.strip()}

只输出 JSON，不要其他文字："""

    try:
        resp = client.chat.completions.create(
            model=settings.llm_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=800,
        )
        raw = resp.choices[0].message.content or "null"
        raw = re.sub(r'<think>[\s\S]*?</think>', '', raw).strip()

        # 提取最外层 JSON 对象或数组（贪婪匹配，避免嵌套截断）
        def _extract_json(s: str):
            for start_char, end_char in [('{', '}'), ('[', ']')]:
                idx = s.find(start_char)
                if idx == -1:
                    continue
                depth = 0
                for i, c in enumerate(s[idx:], idx):
                    if c == start_char:
                        depth += 1
                    elif c == end_char:
                        depth -= 1
                        if depth == 0:
                            return s[idx:i+1]
            if 'null' in s:
                return 'null'
            return None

        json_str = _extract_json(raw)
        if not json_str:
            return None
        parsed = json.loads(json_str)
        if parsed is None:
            return None
        # 新格式：{"logic": "or", "groups": [...]}
        if isinstance(parsed, dict) and 'groups' in parsed:
            return parsed
        # 兼容旧格式：纯列表 → 包装成单组
        if isinstance(parsed, list):
            return {"logic": "or", "groups": [parsed]}
    except Exception as e:
        logger.error(f"[意图解析] LLM 解析失败: {e}")
    return None


def _apply_one_condition(df, cond: dict):
    """对单个条件返回布尔 mask"""
    import pandas as pd

    cols = list(df.columns)

    def _best_col(name: str):
        if name in cols:
            return name
        for c in cols:
            if name in str(c) or str(c) in name:
                return c
        return None

    col_name = _best_col(str(cond.get('col', '')))
    op = cond.get('op', '')
    v1 = str(cond.get('v1', '')).strip()
    v2 = str(cond.get('v2', '')).strip()

    if col_name is None:
        logger.warning(f"[条件执行] 找不到列 {cond.get('col')!r}，跳过")
        return pd.Series([True] * len(df), index=df.index)  # 找不到列时不过滤

    try:
        if op == 'eq':
            return df[col_name].astype(str).str.strip() == v1
        elif op == 'neq':
            return df[col_name].astype(str).str.strip() != v1
        elif op == 'contains':
            return df[col_name].astype(str).str.contains(v1, na=False, regex=False)
        elif op == 'not_contains':
            return ~df[col_name].astype(str).str.contains(v1, na=False, regex=False)
        elif op == 'gt':
            return pd.to_numeric(df[col_name], errors='coerce') > float(v1)
        elif op == 'lt':
            return pd.to_numeric(df[col_name], errors='coerce') < float(v1)
        elif op == 'gte':
            return pd.to_numeric(df[col_name], errors='coerce') >= float(v1)
        elif op == 'lte':
            return pd.to_numeric(df[col_name], errors='coerce') <= float(v1)
        elif op == 'between':
            col_dt = pd.to_datetime(df[col_name], errors='coerce')
            s = pd.to_datetime(v1, errors='coerce')
            e = pd.to_datetime(v2, errors='coerce')
            if pd.notna(s) and pd.notna(e):
                return (col_dt >= s) & (col_dt <= e)
            num = pd.to_numeric(df[col_name], errors='coerce')
            return (num >= float(v1)) & (num <= float(v2))
        elif op == 'empty':
            return df[col_name].isna() | (df[col_name].astype(str).str.strip() == '')
        elif op == 'not_empty':
            return df[col_name].notna() & (df[col_name].astype(str).str.strip() != '')
    except Exception as e:
        logger.warning(f"[条件执行] 条件 {cond} 执行失败: {e}，跳过")

    return pd.Series([True] * len(df), index=df.index)


def _apply_filter_conditions(df, parsed):
    """
    将结构化过滤条件应用到 DataFrame，支持 OR 分组（组内 AND，组间 OR）。

    parsed 格式：
      {"logic": "or", "groups": [[cond, ...], [cond, ...], ...]}
    或兼容旧格式（纯列表，视为单组 AND）。
    """
    import pandas as pd

    # 兼容旧格式
    if isinstance(parsed, list):
        parsed = {"logic": "or", "groups": [parsed]}

    groups = parsed.get("groups", [])
    if not groups:
        return df

    # 每组内部 AND，组间 OR
    or_mask = pd.Series([False] * len(df), index=df.index)
    for group in groups:
        group_mask = pd.Series([True] * len(df), index=df.index)
        for cond in group:
            group_mask = group_mask & _apply_one_condition(df, cond)
        logger.debug(f"[条件执行] 组 {group} → {group_mask.sum()} 行")
        or_mask = or_mask | group_mask

    result = df[or_mask].reset_index(drop=True)
    logger.debug(f"[条件执行] OR合并后 → {len(result)} 行")
    return result


def _filter_dataframe_by_instruction(df, user_instruction: str):
    """
    用自然语言指令对 DataFrame 做行筛选。

    策略：
      1. 规则层：直接解析简单条件（无 LLM，速度最快）
      2. LLM 意图解析：LLM 一次调用解析为结构化条件，规则引擎执行（不逐批处理数据）
      3. 非筛选指令（填表说明等）：直接返回全量数据
    """
    import pandas as pd

    if not user_instruction or not user_instruction.strip():
        return df

    total_rows = len(df)
    if total_rows == 0:
        return df

    # ── 1. 规则层：简单条件直接用 pandas，无需任何 LLM 调用 ──────────
    result = _rule_based_filter(df, user_instruction)
    if result is not None:
        logger.info(f"[行筛选] 规则层命中，{total_rows} → {len(result)} 行，无 LLM 调用")
        return result

    # ── 2. LLM 解析意图（一次调用），规则引擎执行过滤 ─────────────────
    logger.info(f"[行筛选] 规则层未命中，调用 LLM 解析意图（1次）...")
    conditions = _parse_filter_intent(user_instruction, list(df.columns))

    if conditions is None:
        logger.info(f"[行筛选] LLM 判断指令非筛选条件，返回全量数据")
        return df

    groups = conditions.get("groups", []) if isinstance(conditions, dict) else []
    if not groups or all(len(g) == 0 for g in groups):
        logger.info(f"[行筛选] LLM 解析出空条件，返回全量数据")
        return df

    total_conds = sum(len(g) for g in groups)
    logger.info(f"[行筛选] LLM 解析出 {len(groups)} 组 / {total_conds} 个条件: {conditions}")
    result = _apply_filter_conditions(df, conditions)
    logger.info(f"[行筛选] 条件执行完毕，{total_rows} → {len(result)} 行")
    return result


# ────────────────────────────────────────────────────────────
# 工具函数：提取模板表头/字段
# ────────────────────────────────────────────────────────────

def extract_template_fields(template_path: str) -> dict:
    """
    从模板文件中提取需要填写的字段。
    优先识别 {{占位符}}，若无占位符则用 LLM 分析表头。
    返回:
        {
            "fields": ["字段1", "字段2", ...],
            "method": "placeholder" | "llm",
        }
    """
    ext = os.path.splitext(template_path)[1].lower()
    placeholder_pattern = re.compile(r'\{\{(.+?)\}\}')
    found = set()

    if ext == '.docx':
        from docx import Document
        doc = Document(template_path)
        for para in doc.paragraphs:
            found.update(placeholder_pattern.findall(para.text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        found.update(placeholder_pattern.findall(para.text))

    elif ext == '.xlsx':
        from openpyxl import load_workbook
        wb = load_workbook(template_path, data_only=True)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        found.update(placeholder_pattern.findall(cell.value))
        # xlsx 模板：直接用 openpyxl 读取表头列名（绕过 LLM 识别）
        if not found:
            from openpyxl.cell.cell import MergedCell
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    row_vals = []
                    for cell in row:
                        if isinstance(cell, MergedCell):
                            continue
                        v = cell.value
                        if v is None:
                            continue
                        v_str = str(v).strip().strip("'").strip()
                        if not v_str:
                            continue
                        row_vals.append(v_str)
                    if row_vals:
                        found.update(row_vals)
                        break
                if found:
                    break

    if found:
        logger.info(f"占位符检测：发现 {len(found)} 个占位符字段: {found}")
        # xlsx 表格模板用智能追加模式（不是单值占位符替换）
        if ext == '.xlsx':
            return {"fields": list(found), "method": "llm"}
        return {"fields": list(found), "method": "placeholder"}

    # ── 无占位符：用 LLM 分析模板表头 ─────────────────────────
    logger.info("模板无占位符，启动 LLM 智能识别表头字段...")
    raw_text = _read_template_text(template_path, ext)
    if not raw_text.strip():
        return {"fields": [], "method": "llm"}

    client = OpenAI(
        api_key=settings.llm_api_key,
        base_url=settings.llm_base_url,
        default_headers=settings.openai_default_headers,
        timeout=120,
    )
    prompt = f"""以下是一个待填写的表格/文档模板内容。请识别出其中所有需要填写数据的字段名称。
字段可能是表格列标题、表单标签、空白行前的描述文字等。

要求：
1. 只返回字段名称列表，不包含表头说明或注释行
2. 严格返回 JSON 格式：{{"fields": ["字段名1", "字段名2", ...]}}
3. 字段名保留原文，不要翻译或修改

模板内容：
{raw_text[:4000]}

请直接输出 JSON："""

    try:
        resp = client.chat.completions.create(
            model=settings.llm_model,
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            temperature=0,
        )
        result = json.loads(resp.choices[0].message.content)
        fields = result.get("fields", [])
        logger.info(f"LLM 识别到 {len(fields)} 个字段: {fields}")
        return {"fields": fields, "method": "llm"}
    except Exception as e:
        logger.error(f"LLM 识别字段失败: {e}")
        return {"fields": [], "method": "llm"}


def _read_template_text(template_path: str, ext: str) -> str:
    """读取模板文件的纯文本内容"""
    try:
        if ext == '.docx':
            from docx import Document
            doc = Document(template_path)
            lines = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        lines.append(" | ".join(row_texts))
            return "\n".join(lines)

        elif ext == '.xlsx':
            from openpyxl import load_workbook
            wb = load_workbook(template_path, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    row_texts = [str(cell.value).strip().strip("'\"") for cell in row
                                 if cell.value is not None and str(cell.value).strip().strip("'\"")]
                    if row_texts:
                        lines.append(" | ".join(row_texts))
            return "\n".join(lines)
    except Exception as e:
        logger.error(f"读取模板文本失败: {e}")
    return ""


# ────────────────────────────────────────────────────────────
# 工具函数：从数据源文档读取文本
# ────────────────────────────────────────────────────────────

def _read_source_texts(source_file_paths: list) -> str:
    """
    读取所有数据源文件的文本内容，拼接为一个大字符串供 LLM 检索。
    每个文件前加文件名标记方便溯源。
    """
    all_texts = []
    for path in source_file_paths:
        ext = os.path.splitext(path)[1].lower()
        filename = os.path.basename(path)
        try:
            text = ""
            if ext in ('.txt', '.md'):
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()

            elif ext == '.docx':
                from docx import Document
                doc = Document(path)
                lines = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        lines.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_texts = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                        lines.append(" | ".join(row_texts))
                text = "\n".join(lines)

            elif ext == '.xlsx':
                import pandas as pd
                import warnings
                import io
                with warnings.catch_warnings():
                    warnings.simplefilter("ignore")
                    xls = pd.ExcelFile(path)
                    sheet_texts = []
                    for sheet_name in xls.sheet_names:
                        df = pd.read_excel(xls, sheet_name=sheet_name)
                        df = df.fillna("")
                        # 保留列标题行（CSV 第一行），确保 LLM 能识别每列含义
                        # 格式：每行 "列名1: 值1, 列名2: 值2, ..."
                        csv_buffer = io.StringIO()
                        df.to_csv(csv_buffer, index=False, encoding='utf-8')
                        csv_content = csv_buffer.getvalue().strip()
                        sheet_texts.append(f"[Sheet: {sheet_name}]\n{csv_content}")
                    # 拼接所有 sheet，限制总字符数避免 token 爆炸
                    text = "\n\n".join(sheet_texts)
                    MAX_SOURCE = 60000  # Qwen2.5-72B 支持128K token，60000字符≈15000token，完全在范围内
                    if len(text) > MAX_SOURCE:
                        # 在换行符处截断，避免从行中间切断导致数据错位
                        cut = text.rfind('\n', 0, MAX_SOURCE)
                        cut = cut if cut > 0 else MAX_SOURCE
                        text = text[:cut] + "\n...[内容过长已截断]..."

            elif ext == '.pdf':
                try:
                    import pdfplumber
                    with pdfplumber.open(path) as pdf:
                        pages = [p.extract_text() or "" for p in pdf.pages]
                    text = "\n".join(pages)
                except Exception:
                    import fitz
                    doc = fitz.open(path)
                    text = "\n".join(page.get_text() for page in doc)
                    doc.close()

            if text.strip():
                all_texts.append(f"=== 文件：{filename} ===\n{text.strip()}")
                logger.info(f"已读取源文件: {filename}，字符数: {len(text)}")

        except Exception as e:
            logger.warning(f"读取源文件 {filename} 失败: {e}")

    return "\n\n".join(all_texts)


# ────────────────────────────────────────────────────────────
# 工具函数：把 docx 里的表格解析成 DataFrame 列表
# ────────────────────────────────────────────────────────────

def _read_docx_tables_as_dataframes(docx_path: str) -> list:
    """
    用 python-docx 把 Word 文档里的所有表格解析为 pandas DataFrame。
    只保留有表头（首行非空）且行数 >= 2 的表格。
    返回 list[DataFrame]，可能为空（纯文字 docx 无表格）。
    """
    import pandas as pd
    from docx import Document

    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.warning(f"[docx解析] 打开文件失败: {e}")
        return []

    dfs = []
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            row_vals = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            rows_data.append(row_vals)

        if len(rows_data) < 2:
            continue

        # 首行作为表头，去掉全空的列
        headers = rows_data[0]
        data_rows = rows_data[1:]

        # 过滤掉完全重复的相邻列（合并单元格导致的列重复）
        seen_indices = {}
        keep_cols = []
        for i, h in enumerate(headers):
            key = h.strip()
            if key not in seen_indices:
                seen_indices[key] = i
                keep_cols.append(i)

        headers = [headers[i] for i in keep_cols]
        data_rows = [[row[i] if i < len(row) else "" for i in keep_cols] for row in data_rows]

        if not any(h.strip() for h in headers):
            continue  # 全空表头跳过

        df = pd.DataFrame(data_rows, columns=headers)
        df = df.fillna("")
        df.columns = [str(c).strip() for c in df.columns]
        dfs.append(df)
        logger.info(f"[docx表格] 解析到表格: {len(df)}行 × {len(df.columns)}列，表头={list(df.columns)}")

    return dfs


# ────────────────────────────────────────────────────────────
# 核心：LLM 提取字段值（返回列表，支持多行）
# ────────────────────────────────────────────────────────────

def _extract_values_by_llm(template_fields: list, source_file_paths: list, source_text: str,
                           max_rows: int = 5, user_instruction: str = "") -> list:
    """
    从数据源中提取各字段的值，混合策略：
    - xlsx 文件：直接用 pandas 按列名读取（最可靠）
    - docx 文件（含表格）：解析表格为 DataFrame，走与 xlsx 相同的路径
    - 其他文件 / docx 纯文字：用 LLM 从文本中提取（兜底）

    返回格式:
        [{"field_name": "字段名", "values": ["值1", "值2", ...]}, ...]
    """
    if not template_fields:
        return []

    answers = []

    # ── xlsx 文件：直接 pandas 读取，行对齐提取 ──────────────────────────────
    xlsx_sources = [p for p in source_file_paths if os.path.splitext(p)[1].lower() == '.xlsx']
    if xlsx_sources:
        import warnings
        import pandas as pd
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")

        for xlsx_path in xlsx_sources:
            try:
                xls = pd.ExcelFile(xlsx_path)
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    df = df.fillna("")

                    # 清理列名
                    df.columns = [str(c).strip().strip("'\"").strip() for c in df.columns]

                    # ── 通用行筛选（支持任意条件：日期范围、国家、金额等）──
                    if user_instruction and user_instruction.strip():
                        df = _filter_dataframe_by_instruction(df, user_instruction)

                    # ── 行对齐：找出模板字段对应的源列，按行统一提取 ──────
                    # col_mapping: template_field -> source_col_name
                    col_mapping = {}
                    for col_name in df.columns:
                        col_clean = col_name.lower().replace(" ", "").replace("_", "").replace("/", "").replace("\\", "")
                        for tf in template_fields:
                            tf_clean = tf.lower().replace(" ", "").replace("_", "").replace("/", "").replace("\\", "")
                            if col_clean == tf_clean and tf not in col_mapping:
                                col_mapping[tf] = col_name
                                break

                    if not col_mapping:
                        continue

                    # 把 df 限制到 max_rows 行，所有字段共享同一套行索引，保证行对齐
                    df_trimmed = df.head(max_rows)

                    for field_name, col_name in col_mapping.items():
                        col_data = df_trimmed[col_name].astype(str).tolist()
                        col_data = [v for v in col_data if v not in ("", "nan", "NaN", "None")]
                        existing = next((a for a in answers if a["field_name"] == field_name), None)
                        if existing:
                            # 追加，但保持长度一致（用空串补齐已有行）
                            existing["values"].extend(col_data)
                        else:
                            answers.append({"field_name": field_name, "values": col_data})

                    # 行对齐修正：确保所有字段的 values 列表等长（短的用空串补齐）
                    max_len = max((len(a["values"]) for a in answers), default=0)
                    for a in answers:
                        if len(a["values"]) < max_len:
                            a["values"].extend([""] * (max_len - len(a["values"])))

            except Exception as e:
                logger.warning(f"pandas 读取 xlsx 失败 {os.path.basename(xlsx_path)}: {e}")

    # ── docx 文件：优先解析表格走结构化路径，和 xlsx 处理方式一致 ───────────
    docx_sources = [p for p in source_file_paths if os.path.splitext(p)[1].lower() == '.docx']
    docx_with_no_table = []  # 没有表格的 docx，走 LLM 文本兜底
    for docx_path in docx_sources:
        dfs = _read_docx_tables_as_dataframes(docx_path)
        if not dfs:
            docx_with_no_table.append(docx_path)
            continue

        # 合并所有与模板字段有列名交集的表，避免"期中/期末"分两张表时丢弃其中一张
        import pandas as pd
        tf_clean_set = {
            tf.lower().replace(" ", "").replace("_", "").replace("/", "").replace("\\", "")
            for tf in template_fields
        }
        matched_dfs = []
        for d in dfs:
            col_clean_set = {
                c.lower().replace(" ", "").replace("_", "").replace("/", "").replace("\\", "")
                for c in d.columns
            }
            if tf_clean_set & col_clean_set:   # 有交集才合并
                matched_dfs.append(d)

        if not matched_dfs:
            # 无列名交集时兜底：取行数最多的表
            matched_dfs = [max(dfs, key=lambda d: len(d))]

        # 列名相同的多张表纵向合并；列名不同的表各自独立提取后在外部合并
        try:
            main_df = pd.concat(matched_dfs, ignore_index=True)
        except Exception:
            main_df = matched_dfs[0]
        logger.info(f"[docx表格] 合并 {len(matched_dfs)} 张表: 共{len(main_df)}行 × {len(main_df.columns)}列")

        # 行筛选（与 xlsx 完全相同的逻辑）
        if user_instruction and user_instruction.strip():
            main_df = _filter_dataframe_by_instruction(main_df, user_instruction)

        # 列名匹配（精确匹配，与 xlsx 路径相同）
        col_mapping = {}
        for col_name in main_df.columns:
            col_clean = col_name.lower().replace(" ", "").replace("_", "").replace("/", "").replace("\\", "")
            for tf in template_fields:
                tf_clean = tf.lower().replace(" ", "").replace("_", "").replace("/", "").replace("\\", "")
                if col_clean == tf_clean and tf not in col_mapping:
                    col_mapping[tf] = col_name
                    break

        if not col_mapping:
            # 精确匹配失败，留给 LLM 兜底
            docx_with_no_table.append(docx_path)
            continue

        df_trimmed = main_df.head(max_rows)
        for field_name, col_name in col_mapping.items():
            col_data = df_trimmed[col_name].astype(str).tolist()
            col_data = [v for v in col_data if v not in ("", "nan", "NaN", "None")]
            existing = next((a for a in answers if a["field_name"] == field_name), None)
            if existing:
                existing["values"].extend(col_data)
            else:
                answers.append({"field_name": field_name, "values": col_data})

        # 行对齐修正
        max_len = max((len(a["values"]) for a in answers), default=0)
        for a in answers:
            if len(a["values"]) < max_len:
                a["values"].extend([""] * (max_len - len(a["values"])))

        logger.info(f"[docx表格] 结构化提取完成，命中 {len(col_mapping)} 个字段")

    # ── 非 xlsx / 无表格 docx 文件：规则提取优先，LLM 兜底 ──────────────────
    non_structured = [p for p in source_file_paths
                      if os.path.splitext(p)[1].lower() not in ('.xlsx', '.docx')]
    non_structured += docx_with_no_table

    if non_structured:
        # 只用非结构化文件的文本（不包含 xlsx CSV），避免 LLM 重复处理已由 pandas 提取的数据
        non_structured_text = _read_source_texts(non_structured)
        if not non_structured_text.strip():
            non_structured_text = source_text  # 兜底：用全量文本

        # 先尝试规则提取（速度快、准确率高、0次LLM调用）
        rule_answers = _rule_extract_from_text(non_structured_text, template_fields)
        if rule_answers:
            logger.info(f"[规则提取] 成功提取 {len(rule_answers)} 个字段，跳过LLM")
            for ra in rule_answers:
                existing = next((a for a in answers if a["field_name"] == ra["field_name"]), None)
                if existing:
                    new_vals = [v for v in ra["values"] if v not in existing["values"]]
                    existing["values"].extend(new_vals)
                else:
                    answers.append(ra)
        else:
            # 规则提取失败，批量发给 LLM 提取所有行（LLM 自行区分数据段落和描述性文字）
            llm_answers = _extract_all_rows_by_llm(template_fields, non_structured_text, max_rows, user_instruction)
            for la in llm_answers:
                existing = next((a for a in answers if a["field_name"] == la["field_name"]), None)
                if existing:
                    existing["values"].extend(la["values"])
                else:
                    answers.append(la)
            logger.info(f"[批量提取] 完成，合并后字段数: {len(answers)}")

    # 过滤掉"未找到"类占位符，保留空串（用于行对齐），限制每字段最多 max_rows 个值
    for a in answers:
        a["values"] = [v for v in a["values"] if v not in ("(未找到)", "(提取失败)")][:max_rows]

    # 最终行对齐：确保所有字段等长（LLM追加后可能长度不一致）
    if answers:
        max_len = max(len(a["values"]) for a in answers)
        for a in answers:
            if len(a["values"]) < max_len:
                a["values"].extend([""] * (max_len - len(a["values"])))

    logger.info(f"字段提取完成，共 {len(answers)} 个字段有值: " + str([(a["field_name"], len(a["values"])) for a in answers]))
    return answers


def _best_match_col(field: str, col_names: list) -> str | None:
    """
    在 col_names 中找到与 field 最匹配的列名（模糊匹配）。
    """
    field_clean = field.lower().replace(" ", "").replace("_", "").replace("/", "").replace("\\", "")
    best_score = 0
    best_col = None
    for col in col_names:
        col_clean = col.lower().replace(" ", "").replace("_", "").replace("/", "").replace("\\", "")
        if col_clean == field_clean:
            return col
        # 计算相似度：共同字符越多分数越高
        common = sum(1 for c in field_clean if c in col_clean)
        score = common / max(len(field_clean), 1)
        if score > best_score and score > 0.5:
            best_score = score
            best_col = col
    return best_col


def _rule_extract_from_text(source_text: str, template_fields: list) -> list:
    """
    用规则从纯文字文档中提取结构化数据，不依赖 LLM。
    策略：
    1. 检测文本是否为"每行/每段一条记录"的格式（如城市GDP报告）
    2. 尝试解析"字段名：值"或"字段名 值"的模式
    3. 成功则返回字段值列表，失败返回空列表（触发LLM兜底）
    """
    import pandas as pd

    lines = [l.strip() for l in source_text.split('\n') if l.strip()]
    if not lines:
        return []

    # 策略1：检测"字段：值"格式（每行一个字段值对）
    # 例如：城市名：上海\nGDP总量：56708.71亿元
    kv_pattern = re.compile(r'^(.{1,15})[：:]\s*(.+)$')
    kv_records = {}
    for line in lines:
        m = kv_pattern.match(line)
        if m:
            key = m.group(1).strip()
            val = m.group(2).strip()
            # 匹配到模板字段
            for tf in template_fields:
                tf_core = re.sub(r'[（(][^）)]*[）)]', '', tf).strip()
                if tf_core in key or key in tf_core or tf in key or key in tf:
                    if tf not in kv_records:
                        kv_records[tf] = []
                    # 清理单位（亿元、万、元等）
                    val_clean = re.sub(r'[，,（(][^）)]*[）)]', '', val)
                    val_clean = re.sub(r'(亿元|万元|万人|万|元|%|％|个|件|次)$', '', val_clean.strip())
                    val_clean = val_clean.replace(',', '').strip()
                    if val_clean:
                        kv_records[tf].append(val_clean)

    if kv_records and any(len(v) > 0 for v in kv_records.values()):
        logger.info(f"[规则提取] 检测到字段:值格式，提取到 {len(kv_records)} 个字段")
        return [{"field_name": k, "values": v} for k, v in kv_records.items() if v]

    # 策略2：检测数字模式提取（适合城市GDP类报告）
    # 每段包含城市名+多个数值，用正则提取
    # 数值模式：逗号分隔的数字，如 38,731.80 或 38731.80
    num_pattern = re.compile(r'[\d,，]+\.?\d*')

    # 找出哪些字段是数值型（字段名含"GDP""人口""收入""元"等）
    numeric_fields = []
    text_fields = []
    for tf in template_fields:
        if any(kw in tf for kw in ['GDP', 'gdp', '人口', '收入', '元', '数量', '金额', '面积', '比例', '%']):
            numeric_fields.append(tf)
        else:
            text_fields.append(tf)

    # 如果没有明显的数值字段，规则层无法处理
    if not numeric_fields:
        return []

    # 按段落提取（每段一个实体）
    paragraphs = [p.strip() for p in source_text.split('\n') if len(p.strip()) > 10]
    if len(paragraphs) < 2:
        return []

    # 检查段落数量是否合理（至少3段才值得规则提取）
    if len(paragraphs) < 3:
        return []

    logger.debug(f"[规则提取] 尝试数值模式提取，{len(paragraphs)} 个段落，{len(numeric_fields)} 个数值字段")
    # 规则层无法可靠提取数值型段落，返回空让LLM处理
    return []


def _split_by_sentences(text: str) -> list:
    """
    按段落切分文本：每个实体（省/城市/记录）是一整段话，以句号结尾。
    策略：按换行符切段落，每段 = 一个实体的完整描述。
    过滤掉太短的片段（<15字符）。
    """
    paragraphs = [p.strip() for p in text.split('\n') if p.strip() and len(p.strip()) >= 15]
    if paragraphs:
        return paragraphs
    # 兜底：按句号切
    parts = re.split(r'(?<=[。！？])', text)
    return [s.strip() for s in parts if s.strip() and len(s.strip()) >= 15]


def _extract_all_rows_by_llm(template_fields: list, source_text: str,
                              max_rows: int, user_instruction: str = "") -> list:
    """
    按句号切分文本，批量发给 LLM 提取每句话的字段值。
    每个含数据的句子 = 一行；描述性句子由 LLM 自动跳过。
    返回格式：[{"field_name": "字段名", "values": ["值1", "值2", ...]}, ...]
    """
    sentences = _split_by_sentences(source_text)
    if not sentences:
        return []

    # 不截断句子总数，确保所有数据句都被处理
    # （描述性句子会被 LLM 跳过，不占 max_rows 名额）

    client = OpenAI(
        api_key=settings.llm_api_key,
        base_url=settings.llm_base_url,
        default_headers=settings.openai_default_headers,
        timeout=90,
    )
    fields_str = "、".join(template_fields)
    instruction_block = f"\n筛选条件：{user_instruction.strip()}" if user_instruction and user_instruction.strip() else ""

    # 先用关键词粗筛：保留包含任意字段名关键词的句子（方案D：规则粗筛减少LLM负担）
    keywords = set()
    for f in template_fields:
        core = re.sub(r'[（(][^）)]*[）)]', '', f).strip()
        if len(core) >= 2:
            keywords.add(core)
        if len(f) >= 2:
            keywords.add(f)

    # 数字模式：含数字的句子大概率是数据句，也保留
    data_sentences = [s for s in sentences
                      if any(kw in s for kw in keywords) or re.search(r'\d', s)]
    # 如果粗筛后太少，退回全量
    if len(data_sentences) < max_rows // 2:
        data_sentences = sentences
    logger.info(f"[批量提取] 原始 {len(sentences)} 句 → 粗筛后 {len(data_sentences)} 句")

    # 分块，每块 8 句，并发调用 LLM
    # 每段约200-400字，8段≈2400字，避免LLM遗忘中间条目
    BATCH_SIZE = 8
    batches = [data_sentences[i:i+BATCH_SIZE] for i in range(0, len(data_sentences), BATCH_SIZE)]

    def _call_one_batch(args):
        batch_idx, batch = args
        numbered = "\n".join(f"{i+1}. {s}" for i, s in enumerate(batch))
        prompt = f"""从以下句子中提取结构化数据行。

需要提取的字段：{fields_str}{instruction_block}

规则：
1. 每个句子描述一个具体实体（如某省、某城市、某人、某条记录）时，提取为一行
2. 全国性/总结性/背景性描述（如"全国新增..."、"各省均..."、"当日..."）直接跳过，不输出
3. 严禁将字段名本身作为值，只返回实际数据（数字、名称、日期等）
4. 找不到的字段填空字符串 ""
5. 同一个实体只输出一行，不要重复

返回 JSON 数组（只包含有具体实体数据的行，顺序与句子顺序一致）：
[{{"字段1":"值","字段2":"值"}}, ...]

句子列表：
{numbered}

JSON："""
        resp = client.chat.completions.create(
            model=settings.llm_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=4000,
        )
        raw = resp.choices[0].message.content or "[]"
        raw = re.sub(r'<think>[\s\S]*?</think>', '', raw).strip()

        idx = raw.find('[')
        if idx == -1:
            return batch_idx, []
        depth, end_idx = 0, -1
        for i, c in enumerate(raw[idx:], idx):
            if c == '[':
                depth += 1
            elif c == ']':
                depth -= 1
                if depth == 0:
                    end_idx = i
                    break
        if end_idx == -1:
            return batch_idx, []
        rows = json.loads(raw[idx:end_idx+1])
        return batch_idx, rows if isinstance(rows, list) else []

    # 并发执行所有批次，按批次索引排序保证顺序
    from concurrent.futures import ThreadPoolExecutor, as_completed
    batch_results = [None] * len(batches)
    with ThreadPoolExecutor(max_workers=min(8, len(batches))) as executor:
        futures = {executor.submit(_call_one_batch, (i, b)): i for i, b in enumerate(batches)}
        for future in as_completed(futures):
            try:
                batch_idx, rows = future.result()
                batch_results[batch_idx] = rows
                logger.info(f"[批量提取] 第{batch_idx+1}批完成：{len(rows)}行有数据")
            except Exception as e:
                logger.error(f"[批量提取] 某批失败: {e}")

    # 按批次顺序合并，标准化字段名
    def _norm(name: str) -> str:
        return name.strip().replace('（', '(').replace('）', ')').replace(' ', '')

    all_rows = []
    for rows in batch_results:
        if not rows:
            continue
        for row in rows:
            if not isinstance(row, dict):
                continue
            norm_row = {}
            for k, v in row.items():
                k_norm = _norm(str(k))
                matched = k_norm
                for tf in template_fields:
                    if _norm(tf) == k_norm:
                        matched = tf
                        break
                val = str(v).strip() if v is not None else ""
                if val and val not in ("(未找到)", "(提取失败)", "nan", "None"):
                    norm_row[matched] = val
            if norm_row:
                all_rows.append(norm_row)

    all_rows = all_rows[:max_rows]

    # 过滤掉完全空行（没有任何有效字段值）
    all_rows = [r for r in all_rows if any(v and v.strip() for v in r.values())]

    # 按主键字段去重（取字段最完整的行），避免同一实体被提取多次
    # 主键 = 第一个文本型字段（通常是地区/名称）
    def _norm(name: str) -> str:
        return name.strip().replace('（', '(').replace('）', ')').replace(' ', '')

    template_norm = {_norm(tf): tf for tf in template_fields}

    # 找主键字段（第一个在模板字段中的文本字段）
    key_field = None
    for row in all_rows:
        for fn in row:
            canonical = template_norm.get(_norm(fn), fn)
            if canonical in template_fields:
                key_field = canonical
                break
        if key_field:
            break

    if key_field:
        dedup = {}
        for row in all_rows:
            # 统一字段名
            norm_row = {}
            for fn, val in row.items():
                canonical = template_norm.get(_norm(fn), fn)
                norm_row[canonical] = val
            key_val = norm_row.get(key_field, "")
            if not key_val:
                # 没有主键值的行直接保留（不参与去重）
                dedup[id(row)] = norm_row
                continue
            filled = sum(1 for v in norm_row.values() if v and v.strip())
            if key_val not in dedup or filled > sum(1 for v in dedup[key_val].values() if v and v.strip()):
                dedup[key_val] = norm_row
        all_rows = list(dedup.values())
        logger.info(f"[批量提取] 去重后 {len(all_rows)} 行")

    logger.info(f"[批量提取] 共提取 {len(all_rows)} 行数据")

    if not all_rows:
        return []

    # 转换为字段数组格式（行对齐）
    all_field_names = []
    for row in all_rows:
        for fn in row:
            fn_norm = _norm(fn)
            canonical = template_norm.get(fn_norm, fn)  # 映射回模板字段名
            if canonical not in all_field_names:
                all_field_names.append(canonical)

    # 重写 all_rows，统一字段名为模板字段名
    normalized_rows = []
    for row in all_rows:
        new_row = {}
        for fn, val in row.items():
            fn_norm = _norm(fn)
            canonical = template_norm.get(fn_norm, fn)
            if canonical not in new_row or not new_row[canonical]:
                new_row[canonical] = val
        normalized_rows.append(new_row)

    result = []
    for fn in all_field_names:
        vals = [row.get(fn, "") for row in normalized_rows]
        result.append({"field_name": fn, "values": vals})

    return result


def _filter_relevant_paragraphs(source_text: str, template_fields: list, max_chars: int = 10000) -> str:
    """
    从源文本中过滤出与模板字段相关的段落，减少传给 LLM 的 token 量。
    策略：
    1. 按段落分割文本
    2. 每个段落检查是否包含任意字段名的关键词
    3. 优先返回相关段落，不足时补充非相关段落
    4. 总长度限制在 max_chars 以内
    """
    # 提取字段名的核心关键词（去掉括号内容和单位）
    keywords = set()
    for f in template_fields:
        # 去掉括号及内容，取核心词
        core = re.sub(r'[（(][^）)]*[）)]', '', f).strip()
        if len(core) >= 2:
            keywords.add(core)
        # 也加入原始字段名
        if len(f) >= 2:
            keywords.add(f)

    paragraphs = [p.strip() for p in source_text.split('\n') if p.strip()]

    relevant = []
    others = []
    for para in paragraphs:
        if any(kw in para for kw in keywords):
            relevant.append(para)
        else:
            others.append(para)

    # 先放相关段落，再补其他段落，总长度不超过 max_chars
    result_parts = []
    total = 0
    for para in relevant + others:
        if total + len(para) + 1 > max_chars:
            break
        result_parts.append(para)
        total += len(para) + 1

    result = '\n'.join(result_parts)
    logger.info(f"[段落过滤] 原始 {len(source_text)} 字符 → 过滤后 {len(result)} 字符，相关段落 {len(relevant)} 个")
    return result


def _extract_by_llm_from_text(template_fields: list, source_text: str, max_rows: int,
                              user_instruction: str = "") -> list:
    """
    用小模型从单个文本段落中提取字段值（并发场景下每段独立调用）。
    """
    # 使用小模型专门做提取，速度快、token消耗少
    EXTRACT_MODEL = "Qwen/Qwen2.5-7B-Instruct"
    client = OpenAI(
        api_key=settings.llm_api_key,
        base_url=settings.llm_base_url,
        default_headers=settings.openai_default_headers,
        timeout=30,  # 小模型+短文本，30s足够
    )

    MAX_SOURCE = 800  # 每条记录最多800字
    if len(source_text) > MAX_SOURCE:
        cut = source_text.rfind('\n', 0, MAX_SOURCE)
        cut = cut if cut > 0 else MAX_SOURCE
        source_text = source_text[:cut]

    fields_str = "、".join(template_fields)

    instruction_block = ""
    if user_instruction and user_instruction.strip():
        instruction_block = f"筛选条件：{user_instruction.strip()}\n"

    prompt = f"""从以下文本中提取一条记录的数据。

字段列表：{fields_str}
{instruction_block}
规则：
1. 每个字段只返回一个值（字符串，不是数组）
2. 严禁将字段名本身作为值，只返回实际数据（数字、名称、日期等）
3. 找不到的字段不要包含在返回结果中

返回JSON格式：
{{"answers":[{{"field_name":"字段名","value":"实际值"}}]}}

文本：
{source_text}

JSON："""

    try:
        resp = client.chat.completions.create(
            model=EXTRACT_MODEL,
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            temperature=0,
            max_tokens=2000,
        )
        raw = resp.choices[0].message.content
        logger.info(f"LLM 原始返回: {raw[:500]}")

        # 健壮解析：兼容 {"answers":[...]} 和纯列表 [...] 两种格式
        try:
            result = json.loads(raw)
        except Exception:
            # JSON 解析失败，尝试提取最外层 JSON
            def _extract_json(s):
                for sc, ec in [('{', '}'), ('[', ']')]:
                    idx = s.find(sc)
                    if idx == -1:
                        continue
                    depth = 0
                    for i, c in enumerate(s[idx:], idx):
                        if c == sc:
                            depth += 1
                        elif c == ec:
                            depth -= 1
                            if depth == 0:
                                return s[idx:i+1]
                return None
            json_str = _extract_json(raw)
            if not json_str:
                return []
            result = json.loads(json_str)

        # 兼容纯列表格式
        if isinstance(result, list):
            answers = result
        else:
            answers = result.get("answers", [])

        # 标准化字段名（去掉全角/半角括号差异、多余空格）
        def _normalize_field_name(name: str) -> str:
            return (name.strip()
                    .replace('（', '(').replace('）', ')')
                    .replace(' ', ''))

        # 统一转成 values 列表
        normalized = []
        for a in answers:
            if not isinstance(a, dict):
                continue
            field_name = _normalize_field_name(str(a.get("field_name", "")))
            if not field_name:
                continue
            # 匹配到模板字段名（模糊匹配，避免括号差异导致不匹配）
            matched_field = field_name
            for tf in template_fields:
                if _normalize_field_name(tf) == field_name:
                    matched_field = tf
                    break

            if "values" in a:
                vals = a["values"]
                if not isinstance(vals, list):
                    vals = [str(vals)] if vals else []
                vals = [str(v).strip() for v in vals if v is not None and str(v).strip()]
            elif "value" in a:
                raw_val = str(a["value"]).strip()
                vals = [v.strip() for v in raw_val.split("\n") if v.strip()]
                if not vals:
                    vals = [raw_val] if raw_val else []
            else:
                vals = []
            vals = [v for v in vals if v not in ("(未找到)", "(提取失败)", "", "nan", "None")]
            if vals:
                normalized.append({"field_name": matched_field, "values": vals})

        logger.info(f"LLM 提取完成，共 {len(normalized)} 个字段")
        return normalized

    except Exception as e:
        logger.error(f"LLM 提取字段值失败: {e}")
        return []


# ────────────────────────────────────────────────────────────
# 核心：写入填充值到模板文件
# ────────────────────────────────────────────────────────────

def _write_to_template(template_path: str, answers: list, output_path: str, method: str,
                       source_file_paths: list = None, user_instruction: str = "") -> bool:
    """
    将提取的字段值写入模板文件。
    - placeholder 模式：替换 {{字段名}}（answers 需转换为旧格式）
    - llm 模式：在对应表头列下方按行追加数据
    """
    ext = os.path.splitext(template_path)[1].lower()

    if method == "placeholder":
        # table_filler 用旧格式 {"answers": [{"field_name":..., "value":...}]}
        # 占位符场景每个字段只有一个值，取第一个
        legacy_answers = [
            {"field_name": a["field_name"], "value": a["values"][0] if a["values"] else ""}
            for a in answers
        ]
        from modules.filler.table_filler import fill_table
        return fill_table(template_path, {"answers": legacy_answers}, output_path)

    if ext == '.xlsx':
        return _write_xlsx_smart(template_path, answers, output_path)
    elif ext == '.docx':
        return _write_docx_smart(template_path, answers, output_path,
                                 source_file_paths=source_file_paths,
                                 user_instruction=user_instruction)
    else:
        import shutil
        shutil.copy(template_path, output_path)
        return True


def _write_xlsx_smart(template_path: str, answers: list, output_path: str) -> bool:
    """
    智能填写 Excel：
    - 遍历所有工作表，找到最佳表头行（非空单元格最多的行）
    - 处理合并单元格：只读主格，避免从属格值覆盖
    - 按表头列名模糊匹配字段，逐行写入数据
    """
    from openpyxl import load_workbook
    try:
        wb = load_workbook(template_path)

        # field_name -> [值列表]
        field_map = {a["field_name"]: a["values"] for a in answers}

        # ── 遍历所有工作表，找到最佳表头行 ─────────────────────────
        best_ws = None
        best_header_row = None
        best_header_cells = {}  # col_idx -> 字段名
        best_score = 0

        for ws in wb.worksheets:
            # 收集合并单元格主格集合（只读主格的值）
            merged_master = set()
            for rng in ws.merged_cells.ranges:
                merged_master.add((rng.min_row, rng.min_col))

            # 遍历所有行，找非空单元格最多的行
            for row_idx, row in enumerate(ws.iter_rows(), start=1):
                cells_this_row = []
                for cell in row:
                    if cell.value is None:
                        continue
                    val = str(cell.value).strip().strip("'")
                    if not val:
                        continue
                    # 无合并单元格时，所有单元格都有效；有合并单元格时，只取主格
                    if merged_master and (cell.row, cell.column) not in merged_master:
                        continue
                    cells_this_row.append((cell.column, val))

                # 评分：优先选择非空单元格数最多的行
                # 同时排除只有1-2个单元格的行（可能是标题行而非表头）
                if len(cells_this_row) >= 3 and len(cells_this_row) > best_score:
                    best_score = len(cells_this_row)
                    best_ws = ws
                    best_header_row = row_idx
                    best_header_cells = {col: name for col, name in cells_this_row}

        # 没有找到表头行
        if best_ws is None or best_header_row is None:
            logger.warning("模板中未找到有效表头行，直接追加数据")
            for a in answers:
                for val in a["values"]:
                    wb.active.append([a["field_name"], val])
            wb.save(output_path)
            return True

        logger.info(f"找到表头行: 第{best_header_row}行，列数: {len(best_header_cells)}, 工作表: {best_ws.title}")

        # ── 计算最多需要写几行 ───────────────────────────────────
        max_data_rows = max(
            (len(vals) for vals in field_map.values() if vals),
            default=1
        )
        if max_data_rows == 0:
            logger.warning("所有字段均未提取到值，复制模板原文")
            import shutil
            shutil.copy(template_path, output_path)
            return True

        start_row = best_header_row + 1

        # ── 逐行写入，处理合并单元格 ──────────────────────────────
        # 收集所有合并区域（行范围），避免写入时触发 openpyxl 冲突
        merged_row_ranges = set()
        for rng in best_ws.merged_cells.ranges:
            if rng.min_row >= best_header_row:
                merged_row_ranges.add((rng.min_row, rng.max_row))

        written_cells = set()  # 已写入的 (row, col)

        logger.info(f"字段映射: {field_map}")

        for data_row_offset in range(max_data_rows):
            write_row = start_row + data_row_offset

            # 跳过已合并的行（表头本身可能跨行）
            skip_this_row = False
            for min_r, max_r in merged_row_ranges:
                if min_r <= write_row <= max_r and write_row != min_r:
                    skip_this_row = True
                    break
            if skip_this_row:
                continue

            for col_idx, col_name in best_header_cells.items():
                # 避免重复写入同一单元格
                if (write_row, col_idx) in written_cells:
                    continue

                matched_value = _fuzzy_match_field(col_name, field_map, data_row_offset)
                if matched_value is not None:
                    best_ws.cell(row=write_row, column=col_idx, value=matched_value)
                    written_cells.add((write_row, col_idx))

        wb.save(output_path)
        logger.info(f"XLSX 智能填写完成: {output_path}，工作表={best_ws.title}，"
                    f"表头行={best_header_row}，共写入 {max_data_rows} 行数据")
        return True

    except Exception as e:
        logger.error(f"XLSX 智能填写失败: {e}")
        # 删除写了一半的不完整文件，避免用户下载到空白模板
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except Exception:
                pass
        return False


def _extract_table_conditions(description: str, columns: list) -> dict | None:
    """
    从表格上方的描述文字中提取该表的筛选条件（一次 LLM 调用）。
    返回 {"logic": "or", "groups": [[cond, ...]]} 或 None（无条件）。
    """
    if not description or not description.strip():
        return None
    return _parse_filter_intent(description, columns)


def _parse_per_table_conditions(user_instruction: str, table_count: int, columns: list) -> list:
    """
    从用户提示词中解析每个表的筛选条件（一次 LLM 调用）。
    返回长度为 table_count 的列表，每个元素是该表的条件 dict 或 None。
    """
    client = OpenAI(
        api_key=settings.llm_api_key,
        base_url=settings.llm_base_url,
        default_headers=settings.openai_default_headers,
        timeout=60,
    )
    cols_str = "、".join(columns)
    prompt = f"""用户提示词描述了需要填写的多个表格，每个表格有各自的筛选条件。
请从提示词中提取每个表格的筛选条件，返回 JSON 数组，数组长度为 {table_count}（对应表一到表{table_count}）。

可用列名：{cols_str}

支持的操作符：eq（等于）、neq（不等于）、contains（包含）、not_contains（不包含）、
gt（大于）、lt（小于）、gte（大于等于）、lte（小于等于）、between（区间）

每个表的条件格式：
{{"logic": "and", "groups": [[{{"col": "列名", "op": "操作符", "v1": "值"}}]]}}
如果某个表没有明确条件，该位置填 null。

用户提示词：{user_instruction.strip()}

只输出 JSON 数组（长度={table_count}），不要其他文字："""

    try:
        resp = client.chat.completions.create(
            model=settings.llm_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=1000,
        )
        raw = resp.choices[0].message.content or "[]"
        raw = re.sub(r'<think>[\s\S]*?</think>', '', raw).strip()

        def _extract_json(s):
            idx = s.find('[')
            if idx == -1:
                return None
            depth = 0
            for i, c in enumerate(s[idx:], idx):
                if c == '[':
                    depth += 1
                elif c == ']':
                    depth -= 1
                    if depth == 0:
                        return s[idx:i+1]
            return None

        json_str = _extract_json(raw)
        if not json_str:
            return []
        parsed = json.loads(json_str)
        if not isinstance(parsed, list):
            return []
        # 补齐长度
        while len(parsed) < table_count:
            parsed.append(None)
        logger.info(f"[多表条件] 解析结果: {parsed}")
        return parsed[:table_count]
    except Exception as e:
        logger.error(f"[多表条件] 解析失败: {e}")
        return []


def _apply_filter_conditions_fuzzy_time(df, conditions: dict):
    """
    与 _apply_filter_conditions 相同，但对时间列的 eq 条件做模糊匹配
    （只比较日期+小时部分，忽略秒/毫秒差异）。
    """
    import pandas as pd

    groups = conditions.get("groups", []) if isinstance(conditions, dict) else []
    if not groups:
        return df

    cols = list(df.columns)

    def _best_col(name):
        if name in cols:
            return name
        for c in cols:
            if name in str(c) or str(c) in name:
                return c
        return None

    def _apply_one_fuzzy(df_, cond):
        col_name = _best_col(str(cond.get('col', '')))
        op = cond.get('op', '')
        v1 = str(cond.get('v1', '')).strip()
        v2 = str(cond.get('v2', '')).strip()

        if col_name is None:
            return pd.Series([True] * len(df_), index=df_.index)

        try:
            # 时间列 eq：先尝试精确匹配，失败则用 contains 前缀匹配
            if op == 'eq':
                exact = df_[col_name].astype(str).str.strip() == v1
                if exact.any():
                    return exact
                # 模糊：用 v1 的前16字符（"2025-11-25 09:00"）做前缀匹配
                prefix = v1[:16]
                if len(prefix) >= 10:
                    return df_[col_name].astype(str).str.startswith(prefix)
                return exact
            else:
                return _apply_one_condition(df_, cond)
        except Exception:
            return pd.Series([True] * len(df_), index=df_.index)

    or_mask = pd.Series([False] * len(df), index=df.index)
    for group in groups:
        group_mask = pd.Series([True] * len(df), index=df.index)
        for cond in group:
            group_mask = group_mask & _apply_one_fuzzy(df, cond)
        or_mask = or_mask | group_mask

    return df[or_mask].reset_index(drop=True)


def _get_docx_table_contexts(doc) -> list:
    """
    遍历 docx body 元素，返回每个表格及其上方的描述段落文本。
    返回: [{"table": table_obj, "description": "上方段落文字"}, ...]
    """
    from docx.oxml.ns import qn
    body = doc.element.body
    result = []
    prev_paras = []  # 收集上方段落

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            text = child.text or ''
            # 收集所有 run 文字
            runs_text = ''.join(r.text or '' for r in child.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
            if runs_text.strip():
                prev_paras.append(runs_text.strip())
        elif tag == 'tbl':
            # 找到表格，取上方最近的段落作为描述
            description = '\n'.join(prev_paras[-5:]) if prev_paras else ''
            # 找到对应的 table 对象
            from docx.table import Table
            tbl_obj = Table(child, doc)
            result.append({"table": tbl_obj, "description": description})
            prev_paras = []  # 重置，下一个表格重新收集

    return result


def _fill_one_docx_table(table, field_map: dict, max_rows: int):
    """将 field_map 的数据填入单个 docx 表格"""
    if len(table.rows) == 0:
        return
    header_row = table.rows[0]
    headers = [cell.text.strip() for cell in header_row.cells]

    existing_empty_rows = []
    for r in table.rows[1:]:
        if all(c.text.strip() == '' for c in r.cells):
            existing_empty_rows.append(r)

    for row_idx in range(max_rows):
        if row_idx < len(existing_empty_rows):
            new_row = existing_empty_rows[row_idx]
        else:
            new_row = table.add_row()
        for i, cell in enumerate(new_row.cells):
            if i < len(headers):
                matched = _fuzzy_match_field(headers[i], field_map, row_idx)
                if matched is not None:
                    cell.text = str(matched)


def _write_docx_smart(template_path: str, answers: list, output_path: str,
                      source_file_paths: list = None, user_instruction: str = "") -> bool:
    """
    智能填写 Word：
    - 在字段标签后追加第一个值（单值字段）
    - 在表格表头下方逐行追加多值数据
    - 多表场景：每个表根据上方描述段落的条件，从数据源中独立筛选数据填入
    """
    from docx import Document
    try:
        doc = Document(template_path)
        global_field_map = {a["field_name"]: a["values"] for a in answers}

        # ── 段落：字段标签后填第一个值 ─────────────────────────
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            for field_name, values in global_field_map.items():
                if not values:
                    continue
                if (field_name in text and
                        (text.endswith(field_name) or
                         text.endswith(field_name + "：") or
                         text.endswith(field_name + ":"))):
                    value = values[0]
                    if para.runs:
                        para.runs[-1].text += str(value)
                    else:
                        para.add_run(str(value))

        # ── 表格：多表场景按用户提示词+描述段落独立筛选数据 ──────────
        table_contexts = _get_docx_table_contexts(doc)
        has_multi_table = len(table_contexts) > 1 and source_file_paths

        # 预加载 xlsx 源数据（多表场景复用）
        source_dfs = []
        if has_multi_table:
            import pandas as pd, warnings
            xlsx_sources = [p for p in source_file_paths if os.path.splitext(p)[1].lower() == '.xlsx']
            for xlsx_path in xlsx_sources:
                try:
                    with warnings.catch_warnings():
                        warnings.simplefilter("ignore")
                        xls = pd.ExcelFile(xlsx_path)
                        for sheet_name in xls.sheet_names:
                            df = pd.read_excel(xls, sheet_name=sheet_name)
                            df = df.fillna("")
                            df.columns = [str(c).strip().strip("'\"").strip() for c in df.columns]
                            source_dfs.append(df)
                except Exception as e:
                    logger.warning(f"[多表填写] 预加载 xlsx 失败: {e}")

        # 从用户提示词解析每个表的条件（一次 LLM 调用）
        per_table_conditions = []
        if has_multi_table and source_dfs and user_instruction:
            all_cols = list(source_dfs[0].columns) if source_dfs else []
            per_table_conditions = _parse_per_table_conditions(
                user_instruction, len(table_contexts), all_cols
            )
            logger.info(f"[多表填写] 从提示词解析出 {len(per_table_conditions)} 组条件: {per_table_conditions}")

        for t_idx, ctx in enumerate(table_contexts):
            table = ctx["table"]
            description = ctx["description"]
            logger.info(f"[多表填写] 表格{t_idx+1} 描述段落: {repr(description[:80])}")
            if len(table.rows) == 0:
                continue

            # 确定本表的 field_map
            if has_multi_table and source_dfs:
                # 优先用从提示词解析的条件，其次用描述段落
                conditions = None
                if t_idx < len(per_table_conditions):
                    conditions = per_table_conditions[t_idx]
                if not conditions:
                    all_cols = list(source_dfs[0].columns) if source_dfs else []
                    conditions = _extract_table_conditions(description, all_cols)
                logger.info(f"[多表填写] 表格{t_idx+1} 条件: {conditions}")

                if conditions:
                    filtered_dfs = []
                    for df in source_dfs:
                        filtered = _apply_filter_conditions_fuzzy_time(df, conditions)
                        if len(filtered) > 0:
                            filtered_dfs.append(filtered)

                    if filtered_dfs:
                        import pandas as pd
                        merged = pd.concat(filtered_dfs, ignore_index=True)
                        logger.info(f"[多表填写] 表格{t_idx+1} 筛选后 {len(merged)} 行")
                        headers = [cell.text.strip() for cell in table.rows[0].cells]
                        table_field_map = {}
                        for col in merged.columns:
                            col_clean = col.lower().replace(" ", "").replace("_", "")
                            for h in headers:
                                h_clean = h.lower().replace(" ", "").replace("_", "")
                                if col_clean == h_clean and h not in table_field_map:
                                    vals = merged[col].astype(str).tolist()
                                    vals = [v for v in vals if v not in ("", "nan", "NaN", "None")]
                                    table_field_map[h] = vals
                        if table_field_map:
                            max_r = max((len(v) for v in table_field_map.values()), default=0)
                            _fill_one_docx_table(table, table_field_map, max_r)
                            continue

                logger.info(f"[多表填写] 表格{t_idx+1} 条件解析失败或无结果，降级使用全局 field_map")

            # 单表或降级：用全局 field_map
            max_r = max((len(v) for v in global_field_map.values() if v), default=0)
            _fill_one_docx_table(table, global_field_map, max_r)

        doc.save(output_path)
        logger.info(f"DOCX 智能填写完成: {output_path}")
        return True

    except Exception as e:
        logger.error(f"DOCX 智能填写失败: {e}")
        # 删除写了一半的不完整文件，避免用户下载到空白模板
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except Exception:
                pass
        return False


def _fuzzy_match_field(col_name: str, field_map: dict, row_idx: int):
    """
    模糊匹配字段名（忽略大小写、空格、下划线、斜杠），返回对应行的值。

    匹配策略（按优先级）：
    1. 精确匹配（忽略空格后相等）
    2. 包含匹配（清理后列名被字段名包含，或字段名被列名包含）
    3. 子串匹配（任一方包含另一方）

    行索引处理：
    - row_idx < len(values): 返回对应行
    - row_idx >= len(values): 返回最后一个值（重复填充）
    - values 为空或 row_idx < 0: 返回 None
    """
    col_lower = col_name.lower().replace(" ", "").replace("_", "").replace("/", "").replace("\\", "")
    candidates = []

    for field_name, values in field_map.items():
        if not values:
            continue
        field_lower = field_name.lower().replace(" ", "").replace("_", "").replace("/", "").replace("\\", "")

        score = 0
        # 精确匹配（最高优先级）
        if col_lower == field_lower:
            score = 3
        # 列名包含字段名（列标题更具体）
        elif field_lower and col_lower in field_lower:
            score = 2
        # 字段名包含列名（字段定义更具体）
        elif col_lower and field_lower in col_lower:
            score = 1
        else:
            continue

        # 取对应行：
        # - 只有1个值（如"报表日期"这种全局单值字段）→ 重复填充是正确的
        # - 多个值（如"成绩""姓名"等明细字段）→ 超出行数说明该行无数据，填空
        if row_idx < len(values):
            val = values[row_idx]
        elif len(values) == 1:
            val = values[0]   # 单值字段：重复填充
        else:
            continue          # 多值字段：该行无数据，跳过（不填充）

        candidates.append((score, field_name, val))

    if not candidates:
        return None

    # 按优先级排序：分数高的在前；同分时字段名越短越精确（"姓名"比"父亲姓名"短，说明更具体）
    candidates.sort(key=lambda x: (x[0], -len(x[1])), reverse=True)
    best = candidates[0]

    # 如果有多个候选且最高分唯一，直接返回
    if len(candidates) > 1 and candidates[0][0] > candidates[1][0]:
        return best[2]
    # 同分时优先返回精确匹配
    exact = next((c for c in candidates if c[0] == 3), None)
    if exact:
        return exact[2]
    # 同分且无精确匹配：取字段名最短的（子串长度越短越具体）
    return best[2]


# ────────────────────────────────────────────────────────────
# 主入口
# ────────────────────────────────────────────────────────────

def extract_and_fill(
    template_path: str,
    source_file_paths: list,
    output_path: str,
    max_rows: int = 50,
    user_instruction: str = "",
) -> bool:
    """
    智能填表主入口：
    1. 分析模板，识别需要填写的字段（占位符或 LLM 表头识别）
    2. 读取所有数据源文件
    3. 用 LLM 从数据源中提取各字段的值（JSON 数组格式，支持多行）
    4. 将提取的值按行写回模板，保存到 output_path

    Args:
        template_path:      模板文件路径（.docx 或 .xlsx）
        source_file_paths:  数据源文件路径列表
        output_path:        填写完成后的输出路径
        max_rows:           表格类字段最多提取行数（默认50）
        user_instruction:   用户自定义筛选条件，如"只取国家为中国和美国的数据"

    Returns:
        True  成功 / False 失败
    """
    logger.info(f"[智能填表] 开始 | 模板: {os.path.basename(template_path)} | "
                f"数据源: {len(source_file_paths)} 个文件"
                + (f" | 条件: {user_instruction}" if user_instruction else ""))

    try:
        field_info = extract_template_fields(template_path)
        fields = field_info["fields"]
        method = field_info["method"]

        if not fields:
            logger.warning("模板中未识别到任何字段，将直接复制模板文件")
            import shutil
            shutil.copy(template_path, output_path)
            return True

        logger.info(f"识别到 {len(fields)} 个字段（方法: {method}）: {fields}")

        source_text = _read_source_texts(source_file_paths)
        if not source_text.strip() and not any(os.path.splitext(p)[1].lower() == '.xlsx' for p in source_file_paths):
            logger.error("数据源文件读取失败或内容为空")
            return False

        answers = _extract_values_by_llm(fields, source_file_paths, source_text,
                                         max_rows=max_rows, user_instruction=user_instruction)
        if not answers:
            logger.error("字段值提取失败，answers 为空")
            return False

        filled = sum(1 for a in answers if a.get("values"))
        logger.info(f"提取结果：{filled}/{len(answers)} 个字段有数据")

        success = _write_to_template(template_path, answers, output_path, method,
                                     source_file_paths=source_file_paths,
                                     user_instruction=user_instruction)

        if success:
            logger.success(f"[智能填表] 完成 ✅ | 输出: {output_path}")
        else:
            logger.error("[智能填表] 写入模板失败")

        return success

    except Exception as e:
        logger.error(f"[智能填表] 异常: {e}")
        return False


def preview_fill(
    template_path: str,
    source_file_paths: list,
    max_rows: int = 50,
    user_instruction: str = "",
) -> list | None:
    """
    预览填表：只做字段识别和值提取，不生成文件。
    返回格式：[{"field_name": "xxx", "values": ["v1","v2"], "method": "llm"}, ...]
    失败时返回 None。
    """
    logger.info(f"[预览填表] 模板: {os.path.basename(template_path)}, 数据源: {len(source_file_paths)} 个"
                + (f", 条件: {user_instruction}" if user_instruction else ""))
    try:
        field_info = extract_template_fields(template_path)
        fields = field_info["fields"]
        method = field_info["method"]

        if not fields:
            logger.warning("模板中未识别到任何字段")
            return []

        source_text = _read_source_texts(source_file_paths)

        answers = _extract_values_by_llm(fields, source_file_paths, source_text,
                                         max_rows=max_rows, user_instruction=user_instruction)
        # 补充 method 字段供前端展示
        for a in answers:
            a["method"] = method

        logger.info(f"[预览填表] 完成，{len(answers)} 个字段")
        return answers

    except Exception as e:
        logger.error(f"[预览填表] 异常: {e}")
        return None


def write_answers_to_template(
    template_path: str,
    answers: list,
    output_path: str,
    source_file_paths: list = None,
    user_instruction: str = "",
) -> bool:
    """
    将用户确认后的 answers 写入模板文件（预览确认后调用）。
    answers 格式：[{"field_name": "xxx", "values": ["v1","v2"]}, ...]
    """
    logger.info(f"[写入模板] {os.path.basename(template_path)}, {len(answers)} 个字段")
    try:
        field_info = extract_template_fields(template_path)
        method = field_info["method"]

        return _write_to_template(template_path, answers, output_path, method,
                                  source_file_paths=source_file_paths,
                                  user_instruction=user_instruction)
    except Exception as e:
        logger.error(f"[写入模板] 失败: {e}")
        return False