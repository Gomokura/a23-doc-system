"""
文档智能操作交互模块 - Word 文档操作实现
支持：段落编辑、格式调整、内容提取、摘要生成等操作
"""
import os
import re
import json
from typing import Dict, List, Optional, Any, Tuple
from loguru import logger
from openai import OpenAI
from config import settings

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ─────────────────────────────────────────────────────────────────────────────
# 中文数字转换
# ─────────────────────────────────────────────────────────────────────────────
def chinese_to_number(text: str) -> Optional[int]:
    """将中文数字转换为阿拉伯数字"""
    cn_map = {
        '零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
        '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
        '百': 100,
        # 完整的"第X"格式
        '第一': 1, '第二': 2, '第三': 3, '第四': 4, '第五': 5,
        '第六': 6, '第七': 7, '第八': 8, '第九': 9, '第十': 10,
        '第十一': 11, '第十二': 12, '第十三': 13, '第十四': 14, '第十五': 15,
        '第十六': 16, '第十七': 17, '第十八': 18, '第十九': 19, '第二十': 20,
        # 阿拉伯数字前缀
        '第1': 1, '第2': 2, '第3': 3, '第4': 4, '第5': 5,
        '第6': 6, '第7': 7, '第8': 8, '第9': 9, '第10': 10,
    }

    # 直接匹配（处理"第二段"这样的完整输入）
    if text in cn_map:
        return cn_map[text]

    # 匹配 "第X段/节/条" 或 "第X" 格式，提取X部分
    match = re.match(r'^第([一二三四五六七八九十百零\d]+)[段节条章]?$', text)
    if match:
        num_text = match.group(1)
        # 先在映射表中查找
        if num_text in cn_map:
            return cn_map[num_text]
        # 处理纯阿拉伯数字
        try:
            return int(num_text)
        except:
            pass
        # 处理中文数字：十一、十二、二、三十二等
        return _parse_chinese_number(num_text)

    return None


def _parse_chinese_number(text: str) -> Optional[int]:
    """解析中文数字字符串（如"十一"、"二十三"、"第二"、"一百二十五"）"""
    if not text:
        return None

    # 如果以"第"开头，先去掉
    if text.startswith('第'):
        text = text[1:]

    # 去除可能的"段节条章"后缀
    text = re.sub(r'[段节条章]$', '', text)

    # 纯数字直接返回
    if text.isdigit():
        return int(text)

    cn_map = {'零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
              '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10, '百': 100}

    result = 0
    temp = 0
    has_value = False

    i = 0
    while i < len(text):
        ch = text[i]
        if ch in cn_map:
            val = cn_map[ch]
            has_value = True
            if val == 100:  # 百
                if temp == 0:
                    temp = 1
                result = (result + temp) * 100
                temp = 0
            elif val >= 10:  # 十
                if temp == 0:
                    temp = 1
                result += temp * val
                temp = 0
            else:  # 0-9
                temp = temp * 10 + val if temp else val
            i += 1
        elif ch.isdigit():
            temp = temp * 10 + int(ch)
            has_value = True
            i += 1
        else:
            return None

    result += temp
    return result if has_value and result > 0 else None


# ─────────────────────────────────────────────────────────────────────────────
# Word 文档操作类
# ─────────────────────────────────────────────────────────────────────────────
class WordDocumentOperations:
    """
    Word 文档操作类
    提供基于自然语言指令的文档编辑功能
    """

    def __init__(self, file_path: str):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装，请运行: pip install python-docx")

        self.file_path = file_path
        self.doc = Document(file_path)
        self._paragraph_map = None  # 缓存：段落映射列表

    # ─────────────────────────────────────────────────────────────────────────
    # 段落映射（正文 / 标题双链）
    #
    # 用户视角的"第1段"默认指正文段落，不含标题。
    # 标题和正文分开建模后：
    #   - 正文段落操作走 _get_body_paragraph_map() / _find_body_paragraph_by_position()
    #   - 标题操作走 _get_headings_map() / _find_heading_by_position()
    #
    # 正文映射结构（_build_body_paragraph_map）：
    # {
    #   'user_index': 1,        # 用户看到的段落编号（正文不含标题，从1开始）
    #   'doc_index': 3,         # python-docx paragraphs[3] 的真实索引
    #   'element': <lxml.element>,
    #   'text': '段落文本',
    #   'text_preview': '段落文本...',
    #   'style': 'Normal',
    #   'para': <Paragraph>,    # 保留段落对象引用
    # }
    #
    # 标题映射结构（_build_headings_map）：
    # {
    #   'user_index': 1,        # 标题序号（从1开始）
    #   'doc_index': 5,         # python-docx 真实索引
    #   'level': '标题 1',       # 样式名（包含级别信息）
    #   'text': '章节标题',
    #   'para': <Paragraph>,
    # }
    # ─────────────────────────────────────────────────────────────────────────

    def _is_non_body_paragraph(self, para, doc_idx: int, total_paragraphs: int) -> bool:
        """
        基于内容特征判断段落是否为非正文内容（标题、元信息等）。
        当文档没有样式信息时，使用内容特征进行智能判断。
        """
        text = para.text.strip()
        if not text:
            return False

        # 规则1：第一段或第二段通常是标题
        if doc_idx <= 1:
            return True

        # 规则2：以 "——" 或 "—"（破折号）开头的通常是副标题
        if text.strip().startswith(('——', '—', '――')):
            return True

        # 规则3：元信息行（撰稿、摄影、作者等）
        meta_patterns = [
            r'^撰稿[：:：]',
            r'^摄影[：:：]',
            r'^作者[：:：]',
            r'^日期[：:：]',
            r'^时间[：:：]',
            r'^审核[：:：]',
            r'^校对[：:：]',
        ]
        for pattern in meta_patterns:
            if re.match(pattern, text):
                return True

        # 规则4：段落非常短（<15字符）且不在文档末尾，可能是小标题
        if len(text) < 15 and doc_idx < total_paragraphs - 3:
            return True

        # 规则5：检查格式特征（加粗、字号差异）
        # 如果段落有特殊的格式设置，可能是标题
        if para.runs:
            # 检查是否第一个run有加粗且内容较短
            first_run = para.runs[0]
            if first_run.bold and len(text) < 30:
                return True

        return False

    def _build_body_paragraph_map(self) -> List[Dict]:
        """
        构建正文段落映射表（不含标题、元信息等非正文内容）。
        当文档没有样式信息时，使用内容特征智能判断。
        """
        mapping = []
        total = len(self.doc.paragraphs)
        for doc_idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            style_name = para.style.name if para.style else 'Normal'

            # 优先使用样式信息
            base_style_name = ''
            if para.style and para.style.base_style:
                try:
                    base_style_name = para.style.base_style.name
                except:
                    pass

            is_heading_by_style = bool(
                re.search(r'标题|heading|subtitle|subhead|cover', style_name, re.I) or
                re.search(r'标题|heading|subtitle|subhead|cover', base_style_name, re.I) if para.style else False
            )

            # 如果样式没有识别出标题，尝试用内容特征
            is_heading_by_content = self._is_non_body_paragraph(para, doc_idx, total)

            # 只有非标题样式 且 非内容特征识别 才加入正文
            if text and not is_heading_by_style and not is_heading_by_content:
                mapping.append({
                    'user_index': len(mapping) + 1,
                    'doc_index': doc_idx,
                    'element': para._element,
                    'text': text,
                    'text_preview': text[:50] + ('...' if len(text) > 50 else ''),
                    'style': style_name,
                    'para': para,
                })
        return mapping

    def _build_headings_map(self) -> List[Dict]:
        """
        构建标题映射表（仅标题，不含正文段落）。
        使用样式和内容特征双重判断。
        """
        mapping = []
        total = len(self.doc.paragraphs)
        for doc_idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ''

            # 样式判断
            is_heading_by_style = bool(
                re.search(r'标题|heading|subtitle|subhead|cover', style_name, re.I) if para.style else False
            )

            # 内容特征判断（用于样式信息不可用的情况）
            is_heading_by_content = self._is_non_body_paragraph(para, doc_idx, total)

            # 任一条件满足即为标题
            if is_heading_by_style or is_heading_by_content:
                mapping.append({
                    'user_index': len(mapping) + 1,
                    'doc_index': doc_idx,
                    'level': style_name,
                    'text': text,
                    'text_preview': text[:50] + ('...' if len(text) > 50 else ''),
                    'para': para,
                })
        return mapping

    # 缓存键
    _body_paragraph_map: Optional[List[Dict]] = None
    _headings_map: Optional[List[Dict]] = None

    def _get_body_paragraph_map(self, force_refresh: bool = False) -> List[Dict]:
        """获取正文段落映射（带缓存）。force_refresh=True 时强制重建。"""
        if self._body_paragraph_map is None or force_refresh:
            self._body_paragraph_map = self._build_body_paragraph_map()
        return self._body_paragraph_map

    def _get_headings_map(self, force_refresh: bool = False) -> List[Dict]:
        """获取标题映射（带缓存）。force_refresh=True 时强制重建。"""
        if self._headings_map is None or force_refresh:
            self._headings_map = self._build_headings_map()
        return self._headings_map

    def _invalidate_cache(self):
        """文档有改动后清空所有缓存。"""
        self._body_paragraph_map = None
        self._headings_map = None
        self._tables = None

    def _get_tables(self) -> List:
        """获取所有表格（带缓存）"""
        if self._tables is None:
            self._tables = list(self.doc.tables)
        return self._tables

    def _remove_paragraph_element(self, para) -> None:
        """真正从 docx XML 里删除段落。"""
        para_element = para._element
        parent = para_element.getparent()
        if parent is not None:
            parent.remove(para_element)

    def _find_body_paragraph_by_position(self, position: str) -> Tuple[int, Any]:
        """
        根据位置描述找到对应的正文段落（不含标题）。
        使用三层优先级策略：

        L1 固定位置词：最后一段 / 开头 / 末尾
        L2 标准编号：第3段 / 第三段 / 第12段
        L3 内容兜底：包含"关键词"的那段

        Returns:
            (user_index, para_object)  用户视角编号 + 段落对象
        """
        para_map = self._get_body_paragraph_map()
        position = str(position or '').strip()

        if not para_map or not position:
            return -1, None

        # ── L1：固定位置词 ───────────────────────────────────────────────
        if any(w in position for w in ('最后', '结尾', '末尾', '文末')):
            last = para_map[-1]
            return last['user_index'], last['para']

        if any(w in position for w in ('开头', '最前')):
            first = para_map[0]
            return first['user_index'], first['para']

        if position in ('第一段', '第一行'):
            first = para_map[0]
            return first['user_index'], first['para']

        # ── L2：标准编号位置 ──────────────────────────────────────────────
        num = chinese_to_number(position)
        if num is not None and num > 0:
            if num <= len(para_map):
                item = para_map[num - 1]
                return item['user_index'], item['para']
            return -1, None

        # ── L3：内容兜底 ─────────────────────────────────────────────────
        content_match = re.search(
            r'(?:包含|有|写了|提到|关于)[""]?(.+?)[""]?(?:的那?段|的?段落|的那?节)',
            position
        )
        if not content_match:
            content_match = re.search(
                r'[""](.+?)[""]那?段', position
            )
        if content_match:
            keyword = content_match.group(1).strip()
            if keyword:
                candidates = [
                    item for item in para_map
                    if keyword in item['text']
                ]
                if len(candidates) == 1:
                    item = candidates[0]
                    return item['user_index'], item['para']
                elif len(candidates) > 1:
                    return candidates[0]['user_index'], candidates[0]['para']
                return -1, None

        # ── 相对位置：第X段后面/前面 ─────────────────────────────────────
        rel_match = re.search(
            r'第([零一二三四五六七八九十百\d]+)段(后面|前面|之后|之前|后|前)',
            position
        )
        if rel_match:
            base_num = chinese_to_number(rel_match.group(0))
            direction = rel_match.group(2)
            if base_num is not None and 1 <= base_num <= len(para_map):
                target_idx = base_num if direction[0] in ('后', '面') else base_num - 2
                if 0 <= target_idx < len(para_map):
                    item = para_map[target_idx]
                    return item['user_index'], item['para']
                return -1, None

        return -1, None

    def _find_heading_by_position(self, position: str) -> Tuple[int, Any]:
        """
        根据位置描述找到对应的标题。
        支持：第1个标题 / 最后标题 / 包含"关键词"的标题

        Returns:
            (user_index, para_object)
        """
        heading_map = self._get_headings_map()
        position = str(position or '').strip()

        if not heading_map or not position:
            return -1, None

        # L1: "最后标题 / 第一个标题 / 最后一个标题"
        if any(w in position for w in ('最后', '结尾', '末尾', '文末')):
            last = heading_map[-1]
            return last['user_index'], last['para']
        if any(w in position for w in ('开头', '最前', '第一个', '首')):
            first = heading_map[0]
            return first['user_index'], first['para']

        # L2: 标准编号
        num = chinese_to_number(position)
        if num is not None and num > 0:
            if num <= len(heading_map):
                item = heading_map[num - 1]
                return item['user_index'], item['para']
            return -1, None

        # L3: 内容兜底
        content_match = re.search(
            r'(?:包含|有|写了|提到|关于)[""]?(.+?)[""]?(?:标题|的那个)',
            position
        )
        if not content_match:
            content_match = re.search(r'[""](.+?)[""]标题', position)
        if content_match:
            keyword = content_match.group(1).strip()
            if keyword:
                candidates = [item for item in heading_map if keyword in item['text']]
                if len(candidates) == 1:
                    return candidates[0]['user_index'], candidates[0]['para']
                elif len(candidates) > 1:
                    return candidates[0]['user_index'], candidates[0]['para']
                return -1, None

        return -1, None

    def _get_paragraphs(self) -> List:
        """兼容旧接口：返回正文段落对象列表（按 user_index 排序）。"""
        return [item['para'] for item in self._get_body_paragraph_map()]

    # ─────────────────────────────────────────────────────────────────────────
    # 基础操作：读取
    # ─────────────────────────────────────────────────────────────────────────

    def get_all_content(self) -> List[Dict[str, Any]]:
        """获取文档所有内容（带结构信息）"""
        result = []

        # 添加段落信息
        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if text:
                result.append({
                    'type': 'paragraph',
                    'index': idx,
                    'content': text,
                    'style': para.style.name if para.style else 'Normal',
                    'is_heading': bool(re.match(r'^#{1,6}\s', text)) or
                                  bool(re.search(r'标题|heading', para.style.name, re.I)) if para.style else False
                })

        # 添加表格信息
        for idx, table in enumerate(self._get_tables()):
            result.append({
                'type': 'table',
                'index': idx,
                'rows': len(table.rows),
                'cols': len(table.columns) if table.rows else 0,
                'preview': self._table_to_text(table, max_rows=3)
            })

        return result

    def _table_to_text(self, table, max_rows: int = 3) -> str:
        """将表格转换为文本预览"""
        lines = []
        for i, row in enumerate(table.rows):
            if i >= max_rows:
                lines.append("...")
                break
            row_text = " | ".join([cell.text.strip()[:20] for cell in row.cells])
            lines.append(row_text)
        return "\n".join(lines)

    # ─────────────────────────────────────────────────────────────────────────
    # 操作：编辑段落
    # ─────────────────────────────────────────────────────────────────────────

    def edit_paragraph(self, position: str, new_content: str, old_content: str = None) -> Dict[str, Any]:
        """
        编辑指定段落的内容

        Args:
            position: 段落位置描述
            new_content: 新内容
            old_content: 要替换的旧内容（可选）

        Returns:
            操作结果
        """
        # 找到目标段落
        idx, para = self._find_body_paragraph_by_position(position)

        if para is None:
            return {'success': False, 'message': f'未找到"{position}"对应的段落'}

        old_text = para.text

        # 如果指定了旧内容，精确替换
        if old_content:
            if old_content in old_text:
                new_text = old_text.replace(old_content, new_content)
                para.text = new_text
                self.doc.save(self.file_path)
                self._invalidate_cache()
                return {
                    'success': True,
                    'message': f'已替换第{idx}段中的内容',
                    'old_content': old_content,
                    'new_content': new_content
                }
            else:
                return {
                    'success': False,
                    'message': f'未在第{idx}段找到"{old_content}"'
                }
        else:
            # 替换整个段落
            para.text = new_content
            self.doc.save(self.file_path)
            self._invalidate_cache()
            return {
                'success': True,
                'message': f'已修改第{idx}段内容',
                'old_content': old_text[:50] + ('...' if len(old_text) > 50 else ''),
                'new_content': new_content
            }

    def add_paragraph(self, position: str, content: str, style: str = 'Normal') -> Dict[str, Any]:
        """
        添加新段落。
        目前优先保证"开头 / 结尾 / 第X段后"这三类场景行为稳定。
        """
        try:
            position = str(position or '结尾').strip()
            content = str(content or '').strip()
            if not content:
                return {'success': False, 'message': '新增段落内容不能为空'}

            if any(word in position for word in ('结尾', '末尾', '最后', '文末')):
                new_para = self.doc.add_paragraph(content)
            elif any(word in position for w in ('开头', '最前')) or position in ('第一段', '第一节', '第一条'):
                para_map = self._get_body_paragraph_map()
                if not para_map:
                    new_para = self.doc.add_paragraph(content)
                else:
                    first_item = para_map[0]
                    new_para = first_item['para'].insert_paragraph_before(content)
            elif '后' in position:
                idx, ref_para = self._find_body_paragraph_by_position(position)
                if ref_para is None:
                    return {'success': False, 'message': f'未找到"{position}"对应的段落'}
                # 插入到参考段落之后：找到参考段落在 body 中的位置，在其后插入新段落元素
                from docx.oxml import OxmlElement
                from docx.text.paragraph import Paragraph
                body = self.doc.element.body
                ref_idx = list(body).index(ref_para._element)
                new_p = OxmlElement('w:p')
                new_r = OxmlElement('w:r')
                new_t = OxmlElement('w:t')
                new_t.text = content
                new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                new_r.append(new_t)
                new_p.append(new_r)
                body.insert(ref_idx + 1, new_p)
                new_para = Paragraph(new_p, self.doc)
            else:
                new_para = self.doc.add_paragraph(content)

            if style and style != 'Normal':
                try:
                    new_para.style = self.doc.styles[style]
                except Exception:
                    logger.warning(f'未找到样式 {style}，已使用默认样式')

            self.doc.save(self.file_path)
            self._invalidate_cache()

            return {
                'success': True,
                'message': '已添加新段落',
                'content': content,
                'position': position,
            }
        except Exception as e:
            return {'success': False, 'message': f'添加段落失败: {str(e)}'}

    def delete_paragraph(self, position: str) -> Dict[str, Any]:
        """
        删除指定段落
        """
        paras = self._get_paragraphs()
        if not paras:
            return {'success': False, 'message': '文档中没有可删除的正文段落'}

        idx, para = self._find_body_paragraph_by_position(position)

        if para is None:
            return {'success': False, 'message': f'未找到"{position}"对应的正文段落'}

        old_content = para.text
        self._remove_paragraph_element(para)
        self.doc.save(self.file_path)
        self._invalidate_cache()

        return {
            'success': True,
            'message': f'已删除第{idx}段',
            'deleted_content': old_content[:50] + ('...' if len(old_content) > 50 else ''),
            'position': position,
        }

    # ─────────────────────────────────────────────────────────────────────────
    # 操作：格式调整
    # ─────────────────────────────────────────────────────────────────────────

    def format_paragraph(self, position: str, **format_params) -> Dict[str, Any]:
        """
        格式化段落

        Args:
            position: 段落位置
            **format_params: 格式参数（bold, italic, font_size, color, alignment等）

        Returns:
            操作结果
        """
        logger.debug(f"[docx_ops] format_paragraph called with position='{position}', params={format_params}")
        idx, para = self._find_body_paragraph_by_position(position)
        logger.debug(f"[docx_ops] _find_body_paragraph_by_position('{position}') => idx={idx}, para={repr(para.text[:30]) if para else None}")
        if para is None:
            return {'success': False, 'message': f'未找到"{position}"对应的正文段落'}

        changed = []

        # 加粗
        if format_params.get('bold'):
            for run in para.runs:
                run.bold = True
            changed.append('加粗')

        # 倾斜
        if format_params.get('italic'):
            for run in para.runs:
                run.italic = True
            changed.append('倾斜')

        # 字号
        font_size = format_params.get('font_size')
        if font_size:
            for run in para.runs:
                run.font.size = Pt(font_size)
            changed.append(f'字号{font_size}')

        # 颜色
        color = format_params.get('color')
        if color:
            rgb = RGBColor.from_string(color)
            for run in para.runs:
                run.font.color.rgb = rgb
            changed.append(f'颜色#{color}')

        # 对齐
        alignment = format_params.get('alignment')
        if alignment:
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                '两端对齐': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
            changed.append(f'对齐方式:{alignment}')

        # 字体
        font_name = format_params.get('font_name')
        if font_name:
            for run in para.runs:
                run.font.name = font_name
            changed.append(f'字体:{font_name}')

        if changed:
            self.doc.save(self.file_path)
            self._invalidate_cache()
            return {
                'success': True,
                'message': f'已对第{idx}段设置: {", ".join(changed)}',
                'format_changes': changed
            }
        else:
            return {'success': True, 'message': '未指定任何格式修改'}

    # ─────────────────────────────────────────────────────────────────────────
    # 操作：标题操作
    # ─────────────────────────────────────────────────────────────────────────

    def format_heading(self, position: str, **format_params) -> Dict[str, Any]:
        """
        格式化标题（加粗、颜色、字号等）。

        Args:
            position: 标题位置描述（如"第1个标题"、"最后标题"、"包含'引言'的标题"）
            **format_params: 格式参数（与 format_paragraph 一致）

        Returns:
            操作结果
        """
        logger.debug(f"[docx_ops] format_heading called with position='{position}', params={format_params}")
        idx, para = self._find_heading_by_position(position)
        if para is None:
            return {'success': False, 'message': f'未找到"{position}"对应的标题'}

        changed = []

        if format_params.get('bold'):
            for run in para.runs:
                run.bold = True
            changed.append('加粗')

        if format_params.get('italic'):
            for run in para.runs:
                run.italic = True
            changed.append('倾斜')

        font_size = format_params.get('font_size')
        if font_size:
            for run in para.runs:
                run.font.size = Pt(font_size)
            changed.append(f'字号{font_size}')

        color = format_params.get('color')
        if color:
            rgb = RGBColor.from_string(color)
            for run in para.runs:
                run.font.color.rgb = rgb
            changed.append(f'颜色#{color}')

        alignment = format_params.get('alignment')
        if alignment:
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                '两端对齐': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
            changed.append(f'对齐方式:{alignment}')

        font_name = format_params.get('font_name')
        if font_name:
            for run in para.runs:
                run.font.name = font_name
            changed.append(f'字体:{font_name}')

        if changed:
            self.doc.save(self.file_path)
            self._invalidate_cache()
            return {
                'success': True,
                'message': f'已对第{idx}个标题设置: {", ".join(changed)}',
                'format_changes': changed
            }
        else:
            return {'success': True, 'message': '未指定任何格式修改'}

    def edit_heading(self, position: str, new_content: str) -> Dict[str, Any]:
        """
        编辑标题文本内容。

        Args:
            position: 标题位置描述
            new_content: 新标题文本

        Returns:
            操作结果
        """
        idx, para = self._find_heading_by_position(position)
        if para is None:
            return {'success': False, 'message': f'未找到"{position}"对应的标题'}

        old_text = para.text
        para.text = new_content
        self.doc.save(self.file_path)
        self._invalidate_cache()
        return {
            'success': True,
            'message': f'已修改第{idx}个标题',
            'old_content': old_text,
            'new_content': new_content
        }

    def delete_heading(self, position: str) -> Dict[str, Any]:
        """
        删除指定标题。

        Args:
            position: 标题位置描述

        Returns:
            操作结果
        """
        headings_map = self._get_headings_map()
        if not headings_map:
            return {'success': False, 'message': '文档中没有可删除的标题'}

        idx, para = self._find_heading_by_position(position)
        if para is None:
            return {'success': False, 'message': f'未找到"{position}"对应的标题'}

        old_content = para.text
        self._remove_paragraph_element(para)
        self.doc.save(self.file_path)
        self._invalidate_cache()
        return {
            'success': True,
            'message': f'已删除第{idx}个标题',
            'deleted_content': old_content,
            'position': position,
        }

    # ─────────────────────────────────────────────────────────────────────────
    # 操作：内容提取
    # ─────────────────────────────────────────────────────────────────────────

    def extract_content(self, extract_type: str = 'all') -> Dict[str, Any]:
        """
        提取文档内容

        Args:
            extract_type: 提取类型 (all, paragraphs, tables, headings)

        Returns:
            提取的内容
        """
        if extract_type == 'all' or extract_type == 'paragraphs':
            paragraphs = [p.text.strip() for p in self._get_paragraphs() if p.text.strip()]
            if extract_type == 'paragraphs':
                return {'success': True, 'content': paragraphs}

        if extract_type == 'all' or extract_type == 'tables':
            tables = []
            for table in self._get_tables():
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            if extract_type == 'tables':
                return {'success': True, 'content': tables}

        if extract_type == 'all' or extract_type == 'headings':
            headings_map = self._get_headings_map()
            headings = [
                {'level': item['level'], 'content': item['text']}
                for item in headings_map
            ]
            if extract_type == 'headings':
                return {'success': True, 'content': headings}

        # 返回全部内容
        return {
            'success': True,
            'content': {
                'paragraphs': [p.text.strip() for p in self._get_paragraphs()],
                'tables': [
                    [[cell.text.strip() for cell in row.cells] for row in table.rows]
                    for table in self._get_tables()
                ],
                'headings': [
                    {'level': item['level'], 'content': item['text']}
                    for item in self._get_headings_map()
                ]
            }
        }

    # ─────────────────────────────────────────────────────────────────────────
    # 操作：替换文本
    # ─────────────────────────────────────────────────────────────────────────

    def replace_text(self, old_text: str, new_text: str, replace_all: bool = True) -> Dict[str, Any]:
        """
        替换文本

        Args:
            old_text: 要替换的文本
            new_text: 替换为的文本
            replace_all: 是否全部替换

        Returns:
            操作结果
        """
        count = 0

        # 替换段落中的文本
        for para in self.doc.paragraphs:
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)
                count += 1
                if not replace_all:
                    break

        # 替换表格中的文本
        for table in self._get_tables():
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_text)
                        count += 1
                        if not replace_all:
                            break

        if count > 0:
            self.doc.save(self.file_path)
            return {
                'success': True,
                'message': f'已替换 {count} 处',
                'replaced_count': count
            }
        else:
            return {
                'success': False,
                'message': f'未找到"{old_text}"',
                'replaced_count': 0
            }

    # ─────────────────────────────────────────────────────────────────────────
    # 操作：生成摘要
    # ─────────────────────────────────────────────────────────────────────────

    def generate_summary(self, max_length: int = 500) -> Dict[str, Any]:
        """
        使用 LLM 生成文档摘要

        Args:
            max_length: 摘要最大字符数

        Returns:
            生成的摘要
        """
        # 收集文档内容
        paragraphs = [p.text.strip() for p in self._get_paragraphs() if p.text.strip()]
        content = "\n\n".join(paragraphs[:20])  # 取前20段

        if not content.strip():
            return {'success': False, 'message': '文档内容为空'}

        # 调用 LLM 生成摘要
        client = OpenAI(api_key=settings.llm_api_key, base_url=settings.llm_base_url)

        prompt = f"""请阅读以下文档内容，生成一个简洁的摘要。

要求：
1. 摘要长度不超过 {max_length} 字
2. 包含文档的核心主题和关键信息
3. 语言简洁明了

文档内容：
{content[:3000]}

请直接输出摘要："""

        try:
            response = client.chat.completions.create(
                model=settings.llm_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=max_length // 2,
            )

            summary = response.choices[0].message.content.strip()

            return {
                'success': True,
                'summary': summary,
                'source_paragraphs': len(paragraphs)
            }

        except Exception as e:
            logger.error(f"生成摘要失败: {e}")
            return {'success': False, 'message': f'生成摘要失败: {str(e)}'}

    def close(self):
        """关闭文档"""
        self.doc = None
        self._paragraphs = None
        self._tables = None


# ─────────────────────────────────────────────────────────────────────────────
# 快捷函数
# ─────────────────────────────────────────────────────────────────────────────
def word_edit_paragraph(file_path: str, position: str, new_content: str,
                        old_content: str = None) -> Dict[str, Any]:
    """编辑 Word 段落"""
    ops = WordDocumentOperations(file_path)
    try:
        return ops.edit_paragraph(position, new_content, old_content)
    finally:
        ops.close()


def word_format_paragraph(file_path: str, position: str, **format_params) -> Dict[str, Any]:
    """格式化 Word 段落"""
    ops = WordDocumentOperations(file_path)
    try:
        return ops.format_paragraph(position, **format_params)
    finally:
        ops.close()


def word_extract_content(file_path: str, extract_type: str = 'all') -> Dict[str, Any]:
    """提取 Word 内容"""
    ops = WordDocumentOperations(file_path)
    try:
        return ops.extract_content(extract_type)
    finally:
        ops.close()


def word_generate_summary(file_path: str, max_length: int = 500) -> Dict[str, Any]:
    """生成 Word 摘要"""
    ops = WordDocumentOperations(file_path)
    try:
        return ops.generate_summary(max_length)
    finally:
        ops.close()


def word_replace_text(file_path: str, old_text: str, new_text: str,
                     replace_all: bool = True) -> Dict[str, Any]:
    """替换 Word 文本"""
    ops = WordDocumentOperations(file_path)
    try:
        return ops.replace_text(old_text, new_text, replace_all)
    finally:
        ops.close()


def word_format_heading(file_path: str, position: str, **format_params) -> Dict[str, Any]:
    """格式化 Word 标题"""
    ops = WordDocumentOperations(file_path)
    try:
        return ops.format_heading(position, **format_params)
    finally:
        ops.close()


def word_edit_heading(file_path: str, position: str, new_content: str) -> Dict[str, Any]:
    """编辑 Word 标题"""
    ops = WordDocumentOperations(file_path)
    try:
        return ops.edit_heading(position, new_content)
    finally:
        ops.close()


def word_delete_heading(file_path: str, position: str) -> Dict[str, Any]:
    """删除 Word 标题"""
    ops = WordDocumentOperations(file_path)
    try:
        return ops.delete_heading(position)
    finally:
        ops.close()
