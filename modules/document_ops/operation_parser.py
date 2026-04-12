"""
文档智能操作交互模块 - 操作指令解析器
基于自然语言处理技术，将用户指令解析为可执行的操作
支持：Word文档编辑、排版、格式调整、内容提取等操作
"""
import json
import re
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from loguru import logger
from openai import OpenAI
from config import settings


# ─────────────────────────────────────────────────────────────────────────────
# 操作类型枚举
# ─────────────────────────────────────────────────────────────────────────────
class OperationType:
    """支持的操作类型"""
    # Word 文档操作
    EDIT_PARAGRAPH = "edit_paragraph"           # 编辑段落内容
    FORMAT_PARAGRAPH = "format_paragraph"       # 格式化段落
    ADD_PARAGRAPH = "add_paragraph"              # 添加段落
    DELETE_PARAGRAPH = "delete_paragraph"       # 删除段落
    EXTRACT_CONTENT = "extract_content"         # 提取内容
    GENERATE_SUMMARY = "generate_summary"       # 生成摘要
    REPLACE_TEXT = "replace_text"               # 替换文本
    FORMAT_HEADING = "format_heading"           # 格式化标题
    EDIT_HEADING = "edit_heading"               # 编辑标题内容
    DELETE_HEADING = "delete_heading"           # 删除标题

    # Excel 操作
    EDIT_CELL = "edit_cell"                     # 编辑单元格
    FORMAT_CELL = "format_cell"                 # 格式化单元格
    ADD_ROW = "add_row"                         # 添加行
    ADD_COLUMN = "add_column"                    # 添加列
    DELETE_ROW = "delete_row"                   # 删除行
    DELETE_COLUMN = "delete_column"             # 删除列
    EXTRACT_TABLE = "extract_table"             # 提取表格
    CALCULATE = "calculate"                     # 计算

    # 通用操作
    CONVERT_FORMAT = "convert_format"           # 格式转换
    MERGE_DOCUMENTS = "merge_documents"         # 合并文档
    SPLIT_DOCUMENT = "split_document"          # 拆分文档
    UNKNOWN = "unknown"                         # 未知操作


# ─────────────────────────────────────────────────────────────────────────────
# 操作参数结构
# ─────────────────────────────────────────────────────────────────────────────
@dataclass
class Operation:
    """解析后的操作对象"""
    operation_type: str
    confidence: float
    parameters: Dict[str, Any]
    original_instruction: str
    reasoning: str = ""


# ─────────────────────────────────────────────────────────────────────────────
# 操作模式字典 - 用于快速匹配
# ─────────────────────────────────────────────────────────────────────────────
OPERATION_PATTERNS = {
    OperationType.EDIT_PARAGRAPH: [
        r"修改.*内容", r"编辑.*段", r"编辑.*段落", r"把.*改成", r"将.*改为",
        r"编辑.*文字", r"修改.*文字", r"更改.*内容"
    ],
    OperationType.FORMAT_PARAGRAPH: [
        r"设置.*字体", r"修改.*字号", r"改变.*颜色",
        r"加粗.*", r"倾斜.*", r"设置.*样式",
        r".*字体.*", r".*字号.*", r".*颜色.*",
        r"标红", r"标蓝", r"标绿", r"标黑", r"标黄",
        r"标为.*色", r"标成.*色", r"设置.*颜色", r"改成.*颜色", r"改为.*颜色",
        r"把.*标红", r"把.*标蓝", r"把.*标绿", r"把.*标黑",
        r"把.*设为.*色", r"把.*设置成.*色", r"把.*改成.*色",
    ],
    OperationType.ADD_PARAGRAPH: [
        r"添加.*段", r"添加.*段落", r"新增.*内容", r"插入.*文字",
        r"添加.*章节", r"在.*后添加"
    ],
    OperationType.DELETE_PARAGRAPH: [
        r"删除.*段", r"删除.*段落", r"移除.*内容", r"清除.*文字",
        r"删掉.*段", r"把.*删除"
    ],
    OperationType.EXTRACT_CONTENT: [
        r"提取.*内容", r"获取.*信息", r"查找.*数据",
        r"搜索.*内容", r"找出.*信息"
    ],
    OperationType.GENERATE_SUMMARY: [
        r"生成.*摘要", r"总结.*内容", r"概括.*要点",
        r"提炼.*主题", r"提取.*主旨", r"写.*摘要"
    ],
    OperationType.REPLACE_TEXT: [
        r"替换.*为", r"把.*替换成", r"全部替换",
        r"批量替换"
    ],
    OperationType.FORMAT_HEADING: [
        r".*标题.*加粗", r".*标题.*倾斜", r".*标题.*颜色",
        r".*标题.*字号", r".*副标题.*加粗", r".*副标题.*红色",
        r".*一级标题.*加粗", r".*二级标题.*颜色",
        r"把.*标题.*改成.*颜色", r"把.*副标题.*改为"
    ],
    OperationType.EDIT_HEADING: [
        r"修改.*标题", r"编辑.*标题", r"把.*标题.*改成",
        r"更改.*副标题", r"把.*副标题.*改为"
    ],
    OperationType.DELETE_HEADING: [
        r"删除.*标题", r"删除.*副标题", r"移除.*标题",
        r"删掉.*一级标题", r"删掉.*二级标题"
    ],
    OperationType.EDIT_CELL: [
        r"修改.*单元格", r"编辑.*表格.*内容", r"填写.*表格",
        r"把.*改成", r"将.*改为", r"设置.*单元格"
    ],
    OperationType.FORMAT_CELL: [
        r"设置.*单元格.*格式", r"修改.*单元格.*样式"
    ],
    OperationType.ADD_ROW: [
        r"添加.*行", r"新增.*数据行", r"插入.*行"
    ],
    OperationType.EXTRACT_TABLE: [
        r"提取.*表格", r"导出.*数据", r"获取.*表格",
        r"提取.*工作表", r"提取.*数据", r"提取.*内容",
        r"获取.*工作表", r"提取.*工作表"
    ],
    OperationType.CONVERT_FORMAT: [
        r"转换.*格式", r"将.*转为", r"导出.*为",
        r"另存为.*格式"
    ],
}


# ─────────────────────────────────────────────────────────────────────────────
# 意图识别与参数提取
# ─────────────────────────────────────────────────────────────────────────────
class OperationParser:
    """
    操作指令解析器
    使用规则匹配 + LLM 辅助解析
    """

    def __init__(self):
        self.client = None

    def _get_client(self) -> OpenAI:
        if self.client is None:
            self.client = OpenAI(
                api_key=settings.llm_api_key,
                base_url=settings.llm_base_url
            )
        return self.client

    def parse(self, instruction: str, file_type: str = "docx", file_path: str = None) -> Operation:
        """
        解析用户指令

        Args:
            instruction: 自然语言指令
            file_type: 文件类型 (docx/xlsx/txt)
            file_path: 文档路径（可选）。若提供，LLM 解析时会带上文档段落列表，
                       避免瞎猜 position 导致越界问题。
        """
        logger.info(f"[操作解析] 收到指令: {instruction}, file_type={file_type}, file_path={file_path}")

        # 1. 读取文档段落上下文（若 file_path 提供了 docx/xlsx）
        document_context = None
        if file_path and file_type == "docx":
            document_context = self._read_docx_context(file_path)
            logger.debug(f"[操作解析] 文档段落上下文: {document_context}")

        # 2. 规则快速匹配
        matched_type, rule_confidence = self._rule_based_match(instruction, file_type)
        rule_params = self._extract_parameters_by_rules(instruction, matched_type)
        logger.info(f"[操作解析] 规则匹配: type={matched_type}, confidence={rule_confidence:.2f}, params={rule_params}")

        # 3. LLM 辅助解析（带上文档上下文）
        operation_type, llm_confidence, llm_params, reasoning = self._llm_assisted_parse(
            instruction, file_type, document_context=document_context
        )
        logger.info(f"[操作解析] LLM解析: type={operation_type}, confidence={llm_confidence:.2f}, params={llm_params}")

        # 4. 综合决策：winner-take-all → always-merge
        #
        # 核心理念：
        #   - position 永远以规则层为主（规则对"最后一段"、"第X段"等表达最稳定，不会因 LLM 幻觉出错）
        #   - 格式/样式参数以 LLM 为主（规则层只识别"加粗"、"红色"等固定词，LLM 还能识别"标红"等灵活表达）
        #   - 操作类型以高置信度者为主
        #
        # 这样即使 LLM confidence 更高，position 也不会被 LLM 可能的幻觉值覆盖。
        #
        merged_params = {}

        # 位置参数：永远以规则为准
        if rule_params.get('position'):
            merged_params['position'] = rule_params['position']
        elif llm_params.get('position'):
            merged_params['position'] = llm_params['position']

        # 格式参数：LLM 为主，规则层补缺
        # 核心思路：LLM 能理解更灵活的表达（如"标红"、"设蓝色"），但可能格式不标准
        # 规则层能提取标准格式（如 FF0000），但表达有限
        # 所以：优先用 LLM 的，如果 LLM 的格式不规范，再用规则层的补上
        format_keys = ('bold', 'italic', 'font_size', 'color', 'alignment', 'font_name', 'content')
        for key in format_keys:
            llm_val = llm_params.get(key)
            rule_val = rule_params.get(key)
            # 优先用 LLM 的值
            if llm_val is not None:
                merged_params[key] = llm_val
            # LLM 没有或为空，才用规则层的
            elif rule_val is not None:
                merged_params[key] = rule_val

        # 内容参数：编辑/替换场景
        for key in ('new_content', 'old_content', 'replace_all', 'extract_type', 'max_length'):
            if rule_params.get(key) is not None:
                merged_params[key] = rule_params[key]
            elif llm_params.get(key) is not None:
                merged_params[key] = llm_params[key]

        # 操作类型：以置信度高的为主
        #
        # 防御性纠正：当指令明确提到"第X段"但 LLM 误返回标题操作时，强制使用规则层结果。
        # 这是防止 LLM 幻觉（把"第二段"误认为"副标题"）的最后一道防线。
        instruction_has_segment = '段' in instruction
        instruction_has_heading_keyword = any(
            kw in instruction for kw in ('标题', '副标题', '一级标题', '二级标题', '章标题')
        )
        llm_returned_heading_op = operation_type in (
            OperationType.FORMAT_HEADING, OperationType.EDIT_HEADING, OperationType.DELETE_HEADING
        )
        rule_returned_paragraph_op = matched_type in (
            OperationType.FORMAT_PARAGRAPH, OperationType.EDIT_PARAGRAPH, OperationType.DELETE_PARAGRAPH
        )

        logger.info(f"[操作解析] 纠正检查: has_segment={instruction_has_segment}, has_heading_kw={instruction_has_heading_keyword}, llm={operation_type}, rule={matched_type}")

        if instruction_has_segment and not instruction_has_heading_keyword \
                and llm_returned_heading_op and rule_returned_paragraph_op:
            # 强制使用规则层的段落操作，不被 LLM 覆盖
            final_type = matched_type
            final_confidence = rule_confidence
            final_reasoning = f"规则优先（LLM误将'第X段'识别为标题，已纠正）"
        elif llm_confidence >= rule_confidence:
            final_type = operation_type
            final_confidence = llm_confidence
            final_reasoning = reasoning
        else:
            final_type = matched_type
            final_confidence = rule_confidence
            final_reasoning = "规则匹配"

        logger.info(f"[操作解析] 最终决策: type={final_type}, confidence={final_confidence:.2f}, reasoning={final_reasoning}")

        return Operation(
            operation_type=final_type,
            confidence=final_confidence,
            parameters=merged_params,
            original_instruction=instruction,
            reasoning=final_reasoning
        )

    def _read_docx_context(self, file_path: str) -> Optional[Dict]:
        """
        读取 docx 文档上下文，供 LLM 解析时参考。
        正文段落和标题分开统计，与 _build_body_paragraph_map / _build_headings_map 一致。
        """
        try:
            import os
            if not os.path.exists(file_path):
                return None
            from docx import Document
            import re
            doc = Document(file_path)
            paragraphs = []
            headings = []
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                style = p.style.name if p.style else 'Normal'
                is_heading = bool(re.search(r'标题|heading', style, re.I))
                display = text[:60] + ('...' if len(text) > 60 else '')
                item = {'text': display, 'style': style}
                if is_heading:
                    headings.append(item)
                else:
                    item['index'] = len(paragraphs) + 1
                    paragraphs.append(item)

            return {
                'body_paragraphs': paragraphs,
                'total_body_paragraphs': len(paragraphs),
                'headings': headings,
                'total_headings': len(headings),
                'note': (
                    f'文档共有 {len(paragraphs)} 个正文段落（不含标题）、{len(headings)} 个标题。'
                    f'"第N段"指正文段落编号（从1开始）；'
                    f'"第N个标题" / "包含X的标题"指标题。'
                )
            }
        except Exception as e:
            logger.warning(f"[操作解析] 读取文档上下文失败: {e}")
            return None

    def _rule_based_match(self, instruction: str, file_type: str) -> tuple:
        """
        基于规则的意图匹配

        Returns:
            (operation_type, confidence)
        """
        # 中文无大小写之分，.lower() 对汉字无效，直接用原字符串避免 Unicode 规范化问题
        text = instruction.strip()

        # 根据文件类型选择匹配模式
        relevant_patterns = {}
        if file_type == "docx":
            for op_type, patterns in OPERATION_PATTERNS.items():
                if op_type not in [OperationType.EDIT_CELL, OperationType.ADD_ROW,
                                    OperationType.EXTRACT_TABLE, OperationType.ADD_COLUMN]:
                    relevant_patterns[op_type] = patterns
        elif file_type in ["xlsx", "xls"]:
            for op_type, patterns in OPERATION_PATTERNS.items():
                if op_type in [OperationType.EDIT_CELL, OperationType.FORMAT_CELL,
                               OperationType.ADD_ROW, OperationType.ADD_COLUMN,
                               OperationType.EXTRACT_TABLE, OperationType.CALCULATE]:
                    relevant_patterns[op_type] = patterns
        else:
            relevant_patterns = OPERATION_PATTERNS

        best_match = OperationType.UNKNOWN
        best_score = 0.0

        for op_type, patterns in relevant_patterns.items():
            for pattern in patterns:
                if re.search(pattern, text):
                    score = len(pattern) / 20.0  # 模式越长，匹配越精确
                    if score > best_score:
                        best_score = score
                        best_match = op_type

        return best_match, min(best_score * 1.2, 0.95)

    def _extract_parameters_by_rules(self, instruction: str, operation_type: str) -> Dict:
        """
        基于规则提取操作参数。

        位置提取优先级（从高到低）：
        1. 固定位置词：最后一段 / 开头 / 末尾 / 第一段
        2. 标准编号：第3段 / 第三段 / 第12段
        3. 相对位置：第3段后 / 第5段前面
        4. 内容兜底：包含"关键词"的那段

        注意：识别不到位置时不设 position 字段，让执行层返回失败，
        绝对不要默认"第1段"。
        """
        params = {}
        instruction = str(instruction or '').strip()

        # ── 位置提取 ──────────────────────────────────────────────────────

        # 1. 固定位置词
        if re.search(r'(最后一?[段节条章行]|末尾一?[段节条章行]|文末一?[段节条章行]|最后那段|最后那一段|末尾那段|文末那段)', instruction):
            params['position'] = '最后一段'
        elif re.search(r'(开头一?[段节条章行]|第一[段节条章行]|最前面一?[段节条章行]|最前[段节条章行])', instruction):
            params['position'] = '第一段'
        # 2. 标准编号位置
        else:
            para_match = re.search(r'第[一二三四五六七八九十百两零\d]+[段节条章行]', instruction)
            if para_match:
                params['position'] = para_match.group()
            # 3. 相对位置：第3段后面 / 第2段之前
            elif re.search(r'第[一二三四五六七八九十百两零\d]+[段节条章行](后面|前面|之后|之前|后|前)', instruction):
                rel_match = re.search(r'第[一二三四五六七八九十百两零\d]+[段节条章行](后面|前面|之后|之前|后|前)', instruction)
                if rel_match:
                    params['position'] = rel_match.group()
            # 4. 内容兜底位置：包含X的那段 / 有X的段落
            elif re.search(r'(?:包含|有|写了|提到|关于|写着)["""\'](.+?)["""\']的那?段', instruction):
                kw_match = re.search(r'(?:包含|有|写了|提到|关于|写着)["""\'](.+?)["""\']的那?段', instruction)
                params['position_type'] = 'keyword'
                params['keyword'] = kw_match.group(1).strip()
                params['position'] = f'包含"{kw_match.group(1).strip()}"的那段'
            elif re.search(r'(?:包含|有|写了|提到)["""\'](.+?)["""\'](?:的段|段落的?那?段)', instruction):
                kw_match = re.search(r'(?:包含|有|写了|提到)["""\'](.+?)["""\'](?:的段|段落的?那?段)', instruction)
                params['position_type'] = 'keyword'
                params['keyword'] = kw_match.group(1).strip()
                params['position'] = f'包含"{kw_match.group(1).strip()}"的那段'
            # "X那段"（无连词）
            elif re.search(r'["""\'](.+?)["""\'](?:那段|的那段|的段落)', instruction):
                kw_match = re.search(r'["""\'](.+?)["""\'](?:那段|的那段|的段落)', instruction)
                kw = kw_match.group(1).strip()
                if len(kw) >= 2:  # 至少2字才当作关键词
                    params['position_type'] = 'keyword'
                    params['keyword'] = kw
                    params['position'] = f'包含"{kw}"的那段'

        # ── 无明确段落位置时的容错：提取引号内容或关键词作为定位 ─────────────
        # 处理"把人名标红"、"把X标红"这类没有"第X段"但有明确内容的情况
        if not params.get('position'):
            # 尝试多种关键词提取模式
            kw = None

            # 模式1：提取引号内容，如"把人名标红"
            quote_match = re.search(r'把["""\'](.+?)["""\']', instruction)
            if quote_match:
                kw = quote_match.group(1).strip()

            # 模式2：提取"把XXX标红"中"标"字前的内容
            if not kw:
                # 匹配"把某人标红"：关键词在"把"和"标"之间
                before_color = re.search(r'把([^把标设改\n\r]{2,20}?)标[红蓝绿黑白黄]', instruction)
                if before_color:
                    kw = before_color.group(1).strip()

            # 模式3：提取"XXX标红"前的内容（无"把"字）
            if not kw:
                before_color2 = re.search(r'^([^把标设改\n\r]{2,20}?)标[红蓝绿黑白黄]', instruction)
                if before_color2:
                    kw = before_color2.group(1).strip()

            if kw and len(kw) >= 1:
                params['position_type'] = 'keyword'
                params['keyword'] = kw
                params['position'] = f'包含"{kw}"的那段'

        # ── 内容提取 ─────────────────────────────────────────────────────

        # 先统一抽取引号里的内容
        quoted_texts = re.findall(r'[""](.+?)[""]', instruction)
        if quoted_texts:
            if operation_type == OperationType.REPLACE_TEXT and len(quoted_texts) >= 2:
                params['old_content'] = quoted_texts[0]
                params['new_content'] = quoted_texts[1]
            elif operation_type == OperationType.EDIT_PARAGRAPH and len(quoted_texts) >= 2 and ('改成' in instruction or '改为' in instruction):
                params['old_content'] = quoted_texts[0]
                params['new_content'] = quoted_texts[1]
            elif operation_type in (OperationType.EDIT_PARAGRAPH, OperationType.ADD_PARAGRAPH):
                params['new_content'] = quoted_texts[-1]

        # "改成/改为"写法
        if 'new_content' not in params:
            new_content_match = re.search(r'改成[""](.+?)[""]', instruction)
            if new_content_match:
                params['new_content'] = new_content_match.group(1)
            else:
                new_content_match = re.search(r'改为[""](.+?)[""]', instruction)
                if new_content_match:
                    params['new_content'] = new_content_match.group(1)

        # "添加/插入"类指令，优先抽新增内容
        if operation_type == OperationType.ADD_PARAGRAPH and 'content' not in params:
            add_match = re.search(r'(?:添加|新增|插入)(?:一段)?[""](.+?)[""]', instruction)
            if add_match:
                params['content'] = add_match.group(1)
            elif params.get('new_content'):
                params['content'] = params['new_content']

        # ── 格式提取 ─────────────────────────────────────────────────────

        font_size_match = re.search(r'(\d+)[号点]?字', instruction)
        if font_size_match:
            params['font_size'] = int(font_size_match.group(1))

        if '加粗' in instruction or '粗体' in instruction:
            params['bold'] = True
        if '倾斜' in instruction or '斜体' in instruction:
            params['italic'] = True

        color_map = {
            '红色': 'FF0000', '蓝色': '0000FF', '绿色': '00FF00',
            '黑色': '000000', '白色': 'FFFFFF', '黄色': 'FFFF00',
            '红': 'FF0000', '蓝': '0000FF', '绿': '00FF00',
            '黑': '000000', '白': 'FFFFFF', '黄': 'FFFF00',
        }
        # 支持多种颜色表达：标红、标为红色、设为蓝色、改成绿色、改为红等
        # 优先匹配"标/设/改 + 颜色词"的组合
        color_found = None
        for prefix in ('标为', '标成', '标', '设为', '设置成', '设置', '改成', '改为', '为'):
            for color in ('红', '蓝', '绿', '黑', '白', '黄'):
                for suffix in ('色', ''):
                    pattern = prefix + color + suffix
                    if pattern in instruction:
                        color_found = color
                        break
            if color_found:
                break

        if not color_found:
            # 兜底：只要有"X色"或"标X"就提取
            for color in ('红', '蓝', '绿', '黑', '白', '黄'):
                if f'{color}色' in instruction or f'标{color}' in instruction:
                    color_found = color
                    break

        if color_found:
            params['color'] = color_map[color_found]

        return params

    def _llm_assisted_parse(self, instruction: str, file_type: str, document_context: Dict = None) -> tuple:
        """
        LLM 辅助解析

        Args:
            document_context: 文档段落上下文。若提供，LLM 会知道文档有多少段、
                            每段大概是什么内容，避免瞎猜 position。

        Returns:
            (operation_type, confidence, parameters, reasoning)
        """
        # 构建文档上下文描述
        ctx_block = ""
        if document_context:
            ctx_block = f"""
## 文档当前状态（重要！请参考）

文档正文段落和标题分开统计：
- 正文段落：{document_context['total_body_paragraphs']} 个，"第N段"中的N指这些段落（从1开始）
- 标题：{document_context['total_headings']} 个，"第N个标题" / "包含X的标题"指这些

正文段落列表：
"""
            for p in document_context.get('body_paragraphs', []):
                ctx_block += f"  第{p['index']}段 [{p['style']}]: 「{p['text']}」\n"
            if document_context.get('headings'):
                ctx_block += "\n标题列表：\n"
                for i, h in enumerate(document_context['headings'], 1):
                    ctx_block += f"  第{i}个标题 [{h['style']}]: 「{h['text']}」\n"
            ctx_block += """
重要：
- "第3段"指正文段落编号（不含标题）
- "第1个标题" / "包含'关键词'的标题"指标题
- position 不要超出文档实际数量
"""
        else:
            ctx_block = """
## 注意
没有提供文档内容。position 可以用：
  - "第3段" / "最后一段" / "第一段"
  - "包含'关键词'的那段"（内容定位）
  - "第3段后面"（相对位置）
不要盲目猜测超出常规范围的值（建议不超过文档总段数的合理上限）。
"""

        prompt = f"""你是一个文档操作指令解析专家。请分析用户的自然语言指令，判断用户想要执行什么操作。

【重要】用户指令可能包含：
- "标红/标蓝/标绿" = 把文字颜色改成红色/蓝色/绿色
- "把人名标红" = 找到包含"人名"的段落，把那段文字标红
- "第二段标蓝" = 把文档第2段的内容标成蓝色

支持的文档操作类型：
## Word文档操作 (docx)
- format_paragraph: 格式化段落（加粗、颜色、字号等）
  示例："把第一段加粗" → format_paragraph, position="第一段", bold=true
  示例："把第三段标红" → format_paragraph, position="第三段", color="FF0000"
  示例："把人名标蓝" → format_paragraph, position="包含'人名'的那段", color="0000FF"
  示例："第二段设成红色" → format_paragraph, position="第二段", color="FF0000"
  示例："把这段改成绿色" → format_paragraph, position="包含'这段'的那段", color="00FF00"

- edit_paragraph: 编辑段落内容（修改/替换文字）
  示例："把第三段改成新内容" → edit_paragraph, position="第三段", new_content="新内容"

- add_paragraph: 添加新段落
  示例："在开头添加一段" → add_paragraph, position="开头"

- delete_paragraph: 删除段落
  示例："删除最后一段" → delete_paragraph, position="最后一段"

- replace_text: 替换文本
  示例："把所有的北京替换成上海" → replace_text, old_content="北京", new_content="上海"

【颜色格式要求】
返回的 color 必须是十六进制格式，如：
- 红色 → "FF0000"
- 蓝色 → "0000FF"
- 绿色 → "00FF00"
- 黑色 → "000000"
- 黄色 → "FFFF00"

position 参数支持：
1. 标准编号："第3段"、"第三段"、"第12段"
2. 固定位置："最后一段"、"第一段"、"开头"
3. 内容定位："包含'关键词'的那段"

{ctx_block}

请返回JSON格式：
{{
    "operation_type": "操作类型",
    "confidence": 0.0-1.0的置信度,
    "parameters": {{"position": "第3段", "bold": true, "color": "FF0000"}},
    "reasoning": "判断理由"
}}

用户指令：{instruction}
文件类型：{file_type}

请直接输出JSON，不要有其他内容："""

        try:
            client = self._get_client()
            response = client.chat.completions.create(
                model=settings.llm_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0,
                max_tokens=500,
            )

            raw = response.choices[0].message.content or "{}"

            # 清理 JSON 响应
            raw = re.sub(r'^```json\n?|```$', '', raw, flags=re.MULTILINE).strip()
            result = json.loads(raw)

            return (
                result.get("operation_type", OperationType.UNKNOWN),
                float(result.get("confidence", 0.5)),
                result.get("parameters", {}),
                result.get("reasoning", "LLM解析")
            )

        except Exception as e:
            logger.warning(f"LLM辅助解析失败: {e}")
            return OperationType.UNKNOWN, 0.0, {}, "解析失败"


# ─────────────────────────────────────────────────────────────────────────────
# 单例访问
# ─────────────────────────────────────────────────────────────────────────────
_parser_instance: Optional[OperationParser] = None


def get_operation_parser() -> OperationParser:
    """获取解析器单例"""
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = OperationParser()
    return _parser_instance


def parse_operation(instruction: str, file_type: str = "docx", file_path: str = None) -> Operation:
    """
    快捷解析函数。

    Args:
        instruction: 自然语言指令
        file_type: 文件类型 (docx/xlsx/txt)
        file_path: 文档路径（可选）。如果提供，LLM 解析时会拿到文档段落列表，
                   避免瞎猜 position 导致"文档只有 N 段却要删第 12 段"的问题。
    """
    parser = get_operation_parser()
    return parser.parse(instruction, file_type, file_path=file_path)
