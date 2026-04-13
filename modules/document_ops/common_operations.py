"""
文档智能操作交互模块 - 通用操作与执行器
支持：格式转换、文档合并、文档拆分等跨格式操作
"""
import os
import json
import shutil
from typing import Dict, List, Optional, Any
from loguru import logger
from openai import OpenAI
from config import settings

# 文档操作导入
from .operation_parser import OperationType, parse_operation, Operation


# ─────────────────────────────────────────────────────────────────────────────
# 格式转换操作
# ─────────────────────────────────────────────────────────────────────────────
class FormatConverter:
    """
    文档格式转换器
    支持常见文档格式之间的转换
    """
    
    @staticmethod
    def docx_to_pdf(input_path: str, output_path: str) -> Dict[str, Any]:
        """
        将 DOCX 转换为 PDF（需要系统安装 LibreOffice 或 Word）
        """
        try:
            # 方案1: 使用 LibreOffice（需要安装）
            import subprocess
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                 os.path.dirname(output_path), input_path],
                capture_output=True,
                timeout=60
            )
            
            if result.returncode == 0:
                converted = input_path.rsplit('.', 1)[0] + '.pdf'
                if converted != output_path and os.path.exists(converted):
                    shutil.move(converted, output_path)
                return {'success': True, 'output_path': output_path}
            
            return {'success': False, 'message': 'LibreOffice 转换失败，请确保已安装 LibreOffice'}
        
        except FileNotFoundError:
            return {'success': False, 'message': '未找到 LibreOffice，请安装后重试'}
        except Exception as e:
            return {'success': False, 'message': f'转换失败: {str(e)}'}
    
    @staticmethod
    def docx_to_markdown(input_path: str, output_path: str) -> Dict[str, Any]:
        """
        将 DOCX 转换为 Markdown
        """
        try:
            from docx import Document
            
            doc = Document(input_path)
            md_lines = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    md_lines.append('')
                    continue
                
                # 处理标题
                style_name = para.style.name.lower() if para.style else ''
                if 'heading 1' in style_name or '标题 1' in style_name:
                    md_lines.append(f'# {text}')
                elif 'heading 2' in style_name or '标题 2' in style_name:
                    md_lines.append(f'## {text}')
                elif 'heading 3' in style_name or '标题 3' in style_name:
                    md_lines.append(f'### {text}')
                else:
                    md_lines.append(text)
            
            # 处理表格
            for table in doc.tables:
                md_lines.append('')
                for i, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    md_lines.append('| ' + ' | '.join(row_data) + ' |')
                    if i == 0:
                        md_lines.append('|' + '|'.join(['---'] * len(row.cells)) + '|')
            
            # 写入文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(md_lines))
            
            return {'success': True, 'output_path': output_path, 'char_count': len('\n'.join(md_lines))}
        
        except Exception as e:
            return {'success': False, 'message': f'转换失败: {str(e)}'}
    
    @staticmethod
    def markdown_to_docx(input_path: str, output_path: str) -> Dict[str, Any]:
        """
        将 Markdown 转换为 DOCX
        """
        try:
            from docx import Document
            import re
            
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            doc = Document()
            
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph('')
                    continue
                
                # 标题
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                # 表格（简化处理）
                elif line.startswith('|'):
                    continue  # 表格处理较复杂，暂时跳过
                else:
                    doc.add_paragraph(line)
            
            doc.save(output_path)
            return {'success': True, 'output_path': output_path}
        
        except Exception as e:
            return {'success': False, 'message': f'转换失败: {str(e)}'}
    
    @staticmethod
    def xlsx_to_csv(input_path: str, output_path: str, sheet_index: int = 0) -> Dict[str, Any]:
        """
        将 XLSX 转换为 CSV
        """
        try:
            import pandas as pd
            
            df = pd.read_excel(input_path, sheet_name=sheet_index)
            df.to_csv(output_path, index=False, encoding='utf-8-sig')
            
            return {
                'success': True,
                'output_path': output_path,
                'rows': len(df),
                'cols': len(df.columns)
            }
        
        except Exception as e:
            return {'success': False, 'message': f'转换失败: {str(e)}'}
    
    @staticmethod
    def csv_to_xlsx(input_path: str, output_path: str) -> Dict[str, Any]:
        """
        将 CSV 转换为 XLSX
        """
        try:
            import pandas as pd
            
            df = pd.read_csv(input_path)
            df.to_excel(output_path, index=False)
            
            return {
                'success': True,
                'output_path': output_path,
                'rows': len(df),
                'cols': len(df.columns)
            }
        
        except Exception as e:
            return {'success': False, 'message': f'转换失败: {str(e)}'}


# ─────────────────────────────────────────────────────────────────────────────
# 文档合并操作
# ─────────────────────────────────────────────────────────────────────────────
class DocumentMerger:
    """
    文档合并器
    支持将多个文档合并为一个
    """
    
    @staticmethod
    def merge_docx(file_paths: List[str], output_path: str) -> Dict[str, Any]:
        """
        合并多个 DOCX 文件
        """
        try:
            from docx import Document
            
            merged_doc = Document()
            
            for i, file_path in enumerate(file_paths):
                if not os.path.exists(file_path):
                    continue
                
                src_doc = Document(file_path)
                
                # 添加分隔符和标题
                if i > 0:
                    merged_doc.add_page_break()
                merged_doc.add_heading(f'文档 {i + 1}: {os.path.basename(file_path)}', level=2)
                
                # 复制段落
                for para in src_doc.paragraphs:
                    if para.text.strip():
                        new_para = merged_doc.add_paragraph(para.text)
                        if para.style:
                            new_para.style = para.style
                
                # 复制表格
                for table in src_doc.tables:
                    new_table = merged_doc.add_table(rows=0, cols=0)
                    new_table.style = table.style
                    for row in table.rows:
                        new_row = new_table.add_row()
                        for i_cell, cell in enumerate(row.cells):
                            new_row.cells[i_cell].text = cell.text
            
            merged_doc.save(output_path)
            
            return {
                'success': True,
                'output_path': output_path,
                'merged_count': len(file_paths)
            }
        
        except Exception as e:
            return {'success': False, 'message': f'合并失败: {str(e)}'}
    
    @staticmethod
    def merge_xlsx(file_paths: List[str], output_path: str, 
                    merge_method: str = 'vertical') -> Dict[str, Any]:
        """
        合并多个 XLSX 文件
        
        Args:
            file_paths: 文件路径列表
            output_path: 输出路径
            merge_method: 合并方式 (vertical=垂直堆叠, horizontal=水平并排)
        """
        try:
            import pandas as pd
            from openpyxl import Workbook
            
            all_dataframes = []
            
            for file_path in file_paths:
                if not os.path.exists(file_path):
                    continue
                try:
                    df = pd.read_excel(file_path)
                    df['source_file'] = os.path.basename(file_path)
                    all_dataframes.append(df)
                except Exception as e:
                    logger.warning(f'读取 {file_path} 失败: {e}')
            
            if not all_dataframes:
                return {'success': False, 'message': '没有可合并的文件'}
            
            if merge_method == 'vertical':
                merged = pd.concat(all_dataframes, ignore_index=True)
            else:
                # horizontal 需要列名对齐
                merged = pd.concat(all_dataframes, axis=1)
            
            merged.to_excel(output_path, index=False)
            
            return {
                'success': True,
                'output_path': output_path,
                'merged_count': len(all_dataframes),
                'total_rows': len(merged)
            }
        
        except Exception as e:
            return {'success': False, 'message': f'合并失败: {str(e)}'}
    
    @staticmethod
    def merge_markdown(file_paths: List[str], output_path: str) -> Dict[str, Any]:
        """
        合并多个 Markdown 文件
        """
        try:
            merged_content = []
            
            for i, file_path in enumerate(file_paths):
                if not os.path.exists(file_path):
                    continue
                
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                if i > 0:
                    merged_content.append('\n\n---\n\n')  # 分隔符
                
                merged_content.append(f'# {os.path.basename(file_path)}\n\n')
                merged_content.append(content)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(''.join(merged_content))
            
            return {
                'success': True,
                'output_path': output_path,
                'merged_count': len(file_paths)
            }
        
        except Exception as e:
            return {'success': False, 'message': f'合并失败: {str(e)}'}


# ─────────────────────────────────────────────────────────────────────────────
# 文档拆分操作
# ─────────────────────────────────────────────────────────────────────────────
class DocumentSplitter:
    """
    文档拆分器
    支持将一个大文档拆分为多个小文档
    """
    
    @staticmethod
    def split_docx_by_paragraphs(input_path: str, output_dir: str,
                                  paragraphs_per_file: int = 50) -> Dict[str, Any]:
        """
        按段落数拆分 DOCX
        """
        try:
            from docx import Document
            
            doc = Document(input_path)
            paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            
            os.makedirs(output_dir, exist_ok=True)
            
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_files = []
            
            for i in range(0, len(paragraphs), paragraphs_per_file):
                chunk = paragraphs[i:i + paragraphs_per_file]
                new_doc = Document()
                
                for para in chunk:
                    new_para = new_doc.add_paragraph(para.text)
                    if para.style:
                        new_para.style = para.style
                
                output_path = os.path.join(output_dir, f'{base_name}_part{i // paragraphs_per_file + 1}.docx')
                new_doc.save(output_path)
                output_files.append(output_path)
            
            return {
                'success': True,
                'output_dir': output_dir,
                'output_files': output_files,
                'file_count': len(output_files)
            }
        
        except Exception as e:
            return {'success': False, 'message': f'拆分失败: {str(e)}'}
    
    @staticmethod
    def split_xlsx_by_rows(input_path: str, output_dir: str,
                           rows_per_file: int = 1000) -> Dict[str, Any]:
        """
        按行数拆分 XLSX
        """
        try:
            import pandas as pd
            
            df = pd.read_excel(input_path)
            
            os.makedirs(output_dir, exist_ok=True)
            
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_files = []
            
            for i in range(0, len(df), rows_per_file):
                chunk = df[i:i + rows_per_file]
                output_path = os.path.join(output_dir, f'{base_name}_part{i // rows_per_file + 1}.xlsx')
                chunk.to_excel(output_path, index=False)
                output_files.append(output_path)
            
            return {
                'success': True,
                'output_dir': output_dir,
                'output_files': output_files,
                'file_count': len(output_files)
            }
        
        except Exception as e:
            return {'success': False, 'message': f'拆分失败: {str(e)}'}


# ─────────────────────────────────────────────────────────────────────────────
# 操作执行器 - 核心调度器
# ─────────────────────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────────────────────
# 参数归一化
# ─────────────────────────────────────────────────────────────────────────────
def _resolve_position(params: Dict[str, Any]) -> str:
    """
    从 LLM 返回的 parameters 中提取段落位置。
    LLM 可能把位置放在 position / paragraph / paragraph_index / target 等字段。
    任何字段都找不到时返回空字符串（由调用方决定如何处理），
    而不是硬编码回退到 '第1段'，避免误删第一段的 Bug。
    """
    for key in ('position', 'paragraph', 'paragraph_index', 'target', 'location'):
        val = params.get(key)
        if val is not None and str(val).strip():
            return str(val).strip()
    return ''


def _normalize_docx_params(op_type: str, params: Dict[str, Any]) -> Dict[str, Any]:
    """兼容预览/LLM 输出的多种参数结构，转换为 docx 执行器可识别格式。"""
    normalized = dict(params or {})

    if op_type != OperationType.FORMAT_PARAGRAPH:
        return normalized

    paragraphs = normalized.get('paragraphs') or []
    if isinstance(paragraphs, list) and paragraphs:
        first = paragraphs[0] or {}
        if not normalized.get('position') and first.get('index') is not None:
            try:
                normalized['position'] = f"第{int(first['index'])}段"
            except Exception:
                pass
        style = str(first.get('style') or '').upper()
        if style == 'BOLD_RED':
            normalized.setdefault('bold', True)
            normalized.setdefault('color', 'FF0000')
        elif style == 'BOLD':
            normalized.setdefault('bold', True)
        elif style == 'RED':
            normalized.setdefault('color', 'FF0000')

    if not normalized.get('position') and normalized.get('index') is not None:
        try:
            normalized['position'] = f"第{int(normalized['index'])}段"
        except Exception:
            pass

    style = str(normalized.get('style') or '').upper()
    if style == 'BOLD_RED':
        normalized.setdefault('bold', True)
        normalized.setdefault('color', 'FF0000')
    elif style == 'BOLD':
        normalized.setdefault('bold', True)
    elif style == 'RED':
        normalized.setdefault('color', 'FF0000')

    return normalized


class OperationExecutor:
    """
    操作执行器
    根据解析的操作类型，调用相应的操作实现
    """
    
    def __init__(self, file_path: str, file_type: str):
        self.file_path = file_path
        self.file_type = file_type.lower()
        
        # 延迟导入，避免循环依赖
        if self.file_type == 'docx':
            from .docx_operations import WordDocumentOperations
            self.ops = WordDocumentOperations(file_path)
        elif self.file_type in ('xlsx', 'xls'):
            from .xlsx_operations import ExcelDocumentOperations
            self.ops = ExcelDocumentOperations(file_path)
        else:
            self.ops = None
    
    def execute(self, operation: Operation) -> Dict[str, Any]:
        """
        执行操作
        
        Args:
            operation: 解析后的操作对象
            
        Returns:
            执行结果
        """
        op_type = operation.operation_type
        params = operation.parameters
        
        logger.info(f"[执行器] 操作类型: {op_type}, 参数: {params}")
        
        if self.ops is None:
            return {'success': False, 'message': f'不支持的文件类型: {self.file_type}'}
        
        # ── Word 文档操作 ──────────────────────────────────────────────
        if self.file_type == 'docx':
            return self._execute_docx_operation(op_type, params)
        
        # ── Excel 文档操作 ────────────────────────────────────────────
        elif self.file_type in ('xlsx', 'xls'):
            return self._execute_xlsx_operation(op_type, params)
        
        return {'success': False, 'message': '未实现该文件类型的操作'}
    
    def _execute_docx_operation(self, op_type: str, params: Dict) -> Dict[str, Any]:
        """执行 Word 文档操作"""
        params = _normalize_docx_params(op_type, params)
        
        if op_type == OperationType.EDIT_PARAGRAPH:
            pos = _resolve_position(params)
            if not pos:
                return {'success': False, 'message': '未能识别目标段落位置，请重新描述（如"第二段"或"最后一段"）'}
            return self.ops.edit_paragraph(
                position=pos,
                new_content=params.get('new_content', ''),
                old_content=params.get('old_content')
            )

        elif op_type == OperationType.FORMAT_PARAGRAPH:
            pos = _resolve_position(params)
            if not pos:
                return {'success': False, 'message': '未能识别目标段落位置，请重新描述（如"第二段"或"最后一段"）'}
            return self.ops.format_paragraph(
                position=pos,
                bold=params.get('bold'),
                italic=params.get('italic'),
                font_size=params.get('font_size'),
                color=params.get('color'),
                alignment=params.get('alignment'),
                font_name=params.get('font_name')
            )

        elif op_type == OperationType.ADD_PARAGRAPH:
            pos = _resolve_position(params) or '结尾'
            return self.ops.add_paragraph(
                position=pos,
                content=params.get('content', ''),
                style=params.get('style', 'Normal')
            )

        elif op_type == OperationType.DELETE_PARAGRAPH:
            pos = _resolve_position(params)
            if not pos:
                return {'success': False, 'message': '未能识别目标段落位置，请重新描述（如"第二段"或"最后一段"）'}
            return self.ops.delete_paragraph(
                position=pos
            )
        
        elif op_type == OperationType.EXTRACT_CONTENT:
            return self.ops.extract_content(
                extract_type=params.get('extract_type', 'all')
            )
        
        elif op_type == OperationType.GENERATE_SUMMARY:
            return self.ops.generate_summary(
                max_length=params.get('max_length', 500)
            )
        
        elif op_type == OperationType.REPLACE_TEXT:
            return self.ops.replace_text(
                old_text=params.get('old_content', ''),
                new_text=params.get('new_content', ''),
                replace_all=params.get('replace_all', True)
            )
        
        elif op_type == OperationType.CONVERT_FORMAT:
            return FormatConverter.docx_to_markdown(
                self.file_path,
                self.file_path.rsplit('.', 1)[0] + '.md'
            )
        
        return {'success': False, 'message': f'未实现操作: {op_type}'}
    
    def _execute_xlsx_operation(self, op_type: str, params: Dict) -> Dict[str, Any]:
        """执行 Excel 文档操作"""
        
        if op_type == OperationType.EDIT_CELL:
            return self.ops.edit_cell(
                cell_ref=params.get('cell', 'A1'),
                value=params.get('value')
            )
        
        elif op_type == OperationType.FORMAT_CELL:
            return self.ops.format_cell(
                cell_ref=params.get('cell', 'A1'),
                bold=params.get('bold'),
                font_size=params.get('font_size'),
                font_color=params.get('color')
            )
        
        elif op_type == OperationType.ADD_ROW:
            return self.ops.add_row(
                position=params.get('position'),
                values=params.get('values')
            )
        
        elif op_type == OperationType.EXTRACT_TABLE:
            return self.ops.extract_table(
                sheet_name=params.get('sheet_name')
            )
        
        elif op_type == OperationType.CALCULATE:
            return self.ops.calculate(
                formula=params.get('formula', ''),
                cell_ref=params.get('cell')
            )
        
        return {'success': False, 'message': f'未实现操作: {op_type}'}
    
    def close(self):
        """关闭操作对象"""
        if self.ops and hasattr(self.ops, 'close'):
            self.ops.close()


# ─────────────────────────────────────────────────────────────────────────────
# 便捷函数
# ─────────────────────────────────────────────────────────────────────────────
def execute_natural_command(file_path: str, instruction: str) -> Dict[str, Any]:
    """
    执行自然语言命令的便捷函数
    
    Args:
        file_path: 文件路径
        instruction: 自然语言指令
        
    Returns:
        执行结果
    """
    # 确定文件类型
    ext = os.path.splitext(file_path)[1].lower()
    file_type = 'docx' if ext == '.docx' else ('xlsx' if ext in ('.xlsx', '.xls') else 'unknown')
    
    if file_type == 'unknown':
        return {'success': False, 'message': f'不支持的文件类型: {ext}'}
    
    # 解析指令
    operation = parse_operation(instruction, file_type)
    
    # 执行操作
    executor = OperationExecutor(file_path, file_type)
    try:
        result = executor.execute(operation)
        result['operation_type'] = operation.operation_type
        result['confidence'] = operation.confidence
        result['reasoning'] = operation.reasoning
        return result
    finally:
        executor.close()
